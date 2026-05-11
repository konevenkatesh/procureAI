"""
scripts/run_comparative_statement_generator.py

═══════════════════════════════════════════════════════════════════
  Sub-block 7 — ComparativeStatementGenerator
  (the visible demo artifact — evaluation committee report)
═══════════════════════════════════════════════════════════════════
Reads everything Module 3 has built (8 BidderProfile, 240
BidEvaluationFinding, 24 EligibilityMatrix, 3 TenderRanking,
6 BidAnomalyFinding = 91 findings per tender) and emits one
evaluation-committee report per tender (Markdown intermediate +
DOCX final). Every claim traces to a kg_node UUID.

Pipeline (Path A — Markdown + DOCX; PDF deferred to L75):
  1. Per-tender input reads (5 sources)
  2. Effective L1 computation (post-ALB-rejection + post-cartel-flag)
  3. Build report data dict (7 parts populated)
  4. Render Markdown intermediate (template-driven, citation-rich)
  5. Render DOCX via python-docx (headings + tables + footers)
  6. Save both to /tmp/comparative_statements/<tender>.{md,docx}
  7. Compute audit_id = SHA256(sorted finding_node_ids[])[0:16]
  8. Emit ComparativeStatement kg_node with file paths + structured
     summary + 5-layer drilldown references

Report shape (7 parts):
  A. Tender Summary
  B. Bidder Participation Overview (8 bidders with verdict breakdown)
  C. Per-Bidder Detailed Evaluation (8 sections, 10-row criteria tables;
     HARD_BLOCK rows bolded; finding_node_id per row)
  D. Ranking (4-row L1-L4 table + ALB block + L1-L2 gap)
  E. Anomaly Findings (CARTEL_SUSPECT + ALB_CORROBORATION evidence)
  F. Committee Recommendation (3 options + Decision blank;
     effective_L1 surfaced)
  G. Audit Trail (findings_consumed=91, rules_cited, audit_id)

Pure aggregator — no edges. Drilldown via finding_node_ids[] arrays +
file path fields. Mirrors Sub-block 4/5/6 pattern (single batch wrapper,
source_ref idempotency, sentinel snapshot pre/post, RC=2 on drift).
═══════════════════════════════════════════════════════════════════
"""
from __future__ import annotations

import os
import sys
import time
import json
import hashlib
import datetime as _dt
import requests
from collections import defaultdict
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO))

from builder.config import settings


# ── Constants ─────────────────────────────────────────────────────────

TYPOLOGY = "ComparativeStatement"
AGGREGATOR_DOC_ID = "comparative_statement_generator_v1"
SOURCE_REF = "sub_block_7:comparative_statement_generator_v1"
REPORT_TEMPLATE_VERSION = "v1"

ARTIFACT_DIR = Path("/tmp/comparative_statements")

REST = settings.supabase_rest_url
H = {"apikey": settings.supabase_anon_key,
     "Authorization": f"Bearer {settings.supabase_anon_key}"}


# ── REST helpers ──────────────────────────────────────────────────────

def rest_get(path, params=None, range_header=None):
    headers = {**H}
    if range_header:
        headers["Range"] = range_header
    r = requests.get(f"{REST}/rest/v1/{path}", params=params or {},
                     headers=headers, timeout=30)
    r.raise_for_status()
    return r.json()


def rest_post(path, body):
    r = requests.post(
        f"{REST}/rest/v1/{path}", json=body,
        headers={**H, "Content-Type": "application/json",
                 "Prefer": "return=representation"}, timeout=30,
    )
    r.raise_for_status()
    return r.json()


def rest_delete(path, params=None):
    r = requests.delete(f"{REST}/rest/v1/{path}", params=params or {},
                        headers=H, timeout=30)
    r.raise_for_status()


def rest_count(path, params=None):
    r = requests.get(
        f"{REST}/rest/v1/{path}",
        params={**(params or {}),
                "select": "doc_id" if path == "kg_nodes" else "edge_id"},
        headers={**H, "Prefer": "count=exact", "Range": "0-0"},
        timeout=30,
    )
    cr = r.headers.get("Content-Range") or ""
    try:
        return int(cr.split("/")[-1])
    except ValueError:
        return -1


PAGE_SIZE = 100


def fetch_all_by_type(node_type: str, extra_params: dict | None = None) -> list[dict]:
    out: list[dict] = []
    page = 0
    while True:
        params = {
            "select":    "node_id,doc_id,label,properties",
            "node_type": f"eq.{node_type}",
            "order":     "doc_id.asc",
        }
        if extra_params:
            params.update(extra_params)
        rows = rest_get("kg_nodes", params,
                        range_header=f"{page * PAGE_SIZE}-{(page + 1) * PAGE_SIZE - 1}")
        out.extend(rows)
        if len(rows) < PAGE_SIZE:
            break
        page += 1
    return out


def fetch_findings_by_node_ids(node_ids: list[str]) -> list[dict]:
    """Bulk fetch findings via in.(...) for the drilldown layer."""
    if not node_ids:
        return []
    # PostgREST in.() filter — chunk if necessary
    out: list[dict] = []
    CHUNK = 50
    for i in range(0, len(node_ids), CHUNK):
        chunk = node_ids[i:i + CHUNK]
        ids_csv = ",".join(chunk)
        rows = rest_get("kg_nodes", {
            "select":  "node_id,doc_id,label,properties",
            "node_id": f"in.({ids_csv})",
        })
        out.extend(rows)
    return out


# ── Effective L1 computation ─────────────────────────────────────────

def compute_effective_l1(tender_ranking_props: dict,
                         anomaly_findings: list[dict]) -> tuple[dict | None, list[str], list[str]]:
    """Returns (effective_l1_entry, alb_rejected_ids, cartel_referred_ids).
    effective_l1_entry: the ranking[] dict for the first non-skipped
                       bidder, or None if all bidders skipped.
    """
    alb_rejected: set[str] = set()
    if tender_ranking_props.get("alb_action_required"):
        alb_rejected.update(tender_ranking_props.get("alb_candidates") or [])
    cartel_referred: set[str] = set()
    for af in anomaly_findings:
        afp = af.get("properties") or {}
        if afp.get("anomaly_class") == "CARTEL_SUSPECT":
            cartel_referred.update(afp.get("primary_bidder_ids") or [])
    ranking = tender_ranking_props.get("ranking") or []
    for entry in ranking:
        bpid = entry.get("bidder_profile_id")
        if bpid in alb_rejected:
            continue
        if bpid in cartel_referred:
            continue
        return entry, sorted(alb_rejected), sorted(cartel_referred)
    return None, sorted(alb_rejected), sorted(cartel_referred)


# ── audit_id ──────────────────────────────────────────────────────────

def compute_audit_id(finding_node_ids: list[str]) -> str:
    """Deterministic 16-char hex from sorted finding UUIDs."""
    h = hashlib.sha256()
    for nid in sorted(finding_node_ids):
        h.update(nid.encode("utf-8"))
        h.update(b"|")
    return h.hexdigest()[:16]


# ── Markdown rendering ───────────────────────────────────────────────

# Display labels (avoid the underscored verdict strings looking like enum-code)
VERDICT_LABEL = {
    "QUALIFIED":              "QUALIFIED",
    "INELIGIBLE":             "INELIGIBLE",
    "GAP_INSUFFICIENT_DATA":  "GAP (insufficient data)",
    "SKIP_NOT_APPLICABLE":    "SKIP (rule not applicable)",
}
AGGREGATE_LABEL = {
    "QUALIFIED":                     "QUALIFIED",
    "FLAGGED_FOR_COMMITTEE_REVIEW":  "FLAGGED for committee review",
    "MARK_FOR_DOCUMENTATION_REVIEW": "MARKED for documentation review",
    "DISQUALIFIED":                  "DISQUALIFIED",
}


def render_markdown(data: dict) -> str:
    """Render the 7-part evaluation-committee report as Markdown."""
    md: list[str] = []
    A = lambda s: md.append(s)

    t = data["tender"]
    A(f"# Comparative Statement of Bids")
    A("")
    A(f"**Tender:** {t['tender_name']}  ")
    A(f"**NIT No.:** {t['tender_nit_no']}  ")
    A(f"**Estimated Contract Value:** ₹{t['tender_ecv_cr']:.2f} crore  ")
    A(f"**Tender Method:** {t['tender_method']} (lowest bid wins among QUALIFIED)  ")
    A(f"**Report generated:** {data['report_generated_at']}  ")
    A(f"**Audit ID:** `{data['audit_id']}`  ")
    A("")
    A("---")
    A("")

    # === PART A — Tender Summary ===
    A("## Part A — Tender Summary")
    A("")
    A(f"- **Tender ID (internal):** `{t['tender_id']}`")
    A(f"- **Estimated Contract Value (ECV):** ₹{t['tender_ecv_cr']:.2f} crore")
    A(f"- **Tender method:** {t['tender_method']}")
    A("")

    # === PART B — Bidder Participation Overview ===
    A("## Part B — Bidder Participation Overview")
    A("")
    counts = data["participation_counts"]
    A(f"**Total bidders received:** {counts['total']}  ")
    A(f"- QUALIFIED: **{counts['qualified']}**  ")
    A(f"- FLAGGED for committee review: **{counts['flagged']}**  ")
    A(f"- MARKED for documentation review: **{counts['mark_for_doc']}**  ")
    A(f"- DISQUALIFIED: **{counts['disqualified']}**  ")
    A("")
    A("### Excluded bidders (drilldown to EligibilityMatrix node):")
    A("")
    A("| Bidder | Aggregate verdict | Exclusion summary | EligibilityMatrix node_id |")
    A("|---|---|---|---|")
    for ex in data["excluded_summaries"]:
        A(f"| {ex['bidder_name']} | {AGGREGATE_LABEL.get(ex['aggregate_verdict'], ex['aggregate_verdict'])} | "
          f"{ex['exclusion_summary']} | `{ex['eligibility_matrix_node_id']}` |")
    A("")

    # === PART C — Per-Bidder Detailed Evaluation ===
    A("## Part C — Per-Bidder Detailed Evaluation")
    A("")
    for bidder_data in data["per_bidder_evaluations"]:
        prof = bidder_data["profile"]
        em = bidder_data["em_props"]
        A(f"### {prof.get('company_name', prof.get('profile_id'))}")
        A("")
        A(f"- **Contractor class:** {prof.get('contractor_class', '?')}")
        A(f"- **Registration:** {prof.get('registration_certificate_no', '?')} (state: {prof.get('registration_state', '?')})")
        A(f"- **PAN / GSTIN:** {prof.get('pan', '?')} / {prof.get('gstin', '?')}")
        A(f"- **Communication address:** {prof.get('communication_address', '?')}")
        A(f"- **Authorized signatory:** {prof.get('authorized_signatory_name', '?')} ({prof.get('authorized_signatory_role', '?')})")
        A("")
        A(f"**Aggregate verdict:** **{AGGREGATE_LABEL.get(em.get('aggregate_verdict'), em.get('aggregate_verdict'))}**  ")
        A(f"({em.get('count_qualified', 0)} QUALIFIED + "
          f"{em.get('count_ineligible_hard_block', 0)} HARD_BLOCK + "
          f"{em.get('count_ineligible_warning', 0)} WARNING + "
          f"{em.get('count_gap', 0)} GAP) of {em.get('criteria_total', 10)} criteria")
        A("")
        A(f"_{em.get('aggregate_reasoning', '')}_")
        A("")
        A("**10-criterion evaluation:**")
        A("")
        A("| # | Criterion | Verdict | Decision reason | Rule | Finding node_id |")
        A("|---|---|---|---|---|---|")
        for i, f in enumerate(bidder_data["findings"], 1):
            fp = f.get("properties") or {}
            v = fp.get("verdict") or "?"
            cons = fp.get("evaluation_consequence") or ""
            v_label = VERDICT_LABEL.get(v, v)
            # Bold the HARD_BLOCK rows
            row_prefix = "**" if cons == "HARD_BLOCK" and v == "INELIGIBLE" else ""
            row_suffix = "**" if row_prefix else ""
            typology = fp.get("typology_code", "?")
            reason = (fp.get("decision_reason") or "").replace("|", "/")[:160]
            rule = fp.get("rule_id", "?")
            nid = f.get("node_id", "?")
            A(f"| {i} | {row_prefix}{typology}{row_suffix} | "
              f"{row_prefix}{v_label}{row_suffix} | "
              f"{row_prefix}{reason}{row_suffix} | `{rule}` | `{nid}` |")
        A("")

    # === PART D — Ranking ===
    A("## Part D — Ranking of QUALIFIED Bidders")
    A("")
    tr_props = data["tender_ranking_props"]
    A("| Rank | Bidder | Bid amount (₹cr) | Premium % vs ECV | ALB flag | Distance from L1 |")
    A("|---|---|---:|---:|---|---:|")
    for r in tr_props.get("ranking") or []:
        alb = "⚠ YES" if r.get("alb_flag") else "—"
        d_cr = r.get("distance_from_l1_cr", 0.0) or 0.0
        d_pct = r.get("distance_from_l1_pct", 0.0) or 0.0
        delta_str = "—" if r["rank_position"] == "L1" else f"+₹{d_cr:.2f}cr ({d_pct:.2f}%)"
        A(f"| {r['rank_position']} | {r['bidder_name']} | "
          f"{r['bid_amount_cr']:.2f} | {r['premium_pct']:.2f} | {alb} | {delta_str} |")
    A("")
    A("### ALB (Abnormally Low Bid) Detection")
    A("")
    A(f"- **Methodology:** `{tr_props.get('alb_threshold_method', '?')}`")
    A(f"- **Average of qualified bids:** ₹{tr_props.get('average_qualified_bid_cr', 0.0):.2f}cr")
    A(f"- **ALB threshold:** ₹{tr_props.get('alb_threshold_cr', 0.0):.2f}cr "
      f"(average × {tr_props.get('alb_multiplier', 0.80)})")
    A(f"- **ALB candidates:** {tr_props.get('alb_candidates') or '(none)'}")
    A(f"- **Action required on L1:** {'**YES** — L1 is ALB candidate' if tr_props.get('alb_action_required') else 'No'}")
    note = tr_props.get("alb_methodology_note") or ""
    if note:
        A("")
        A(f"_Methodology note: {note[:400]}_")
    A("")
    if tr_props.get("l1_l2_gap_cr") is not None:
        A(f"**L1 → L2 gap:** ₹{tr_props['l1_l2_gap_cr']:.2f}cr "
          f"({tr_props.get('l1_l2_gap_pct', 0.0):.2f}%)")
    A("")

    # === PART E — Anomaly Findings ===
    A("## Part E — Anomaly Findings")
    A("")
    if not data["anomaly_findings"]:
        A("_No anomaly findings emitted by CrossBidAnomalyDetector for this tender._")
    for af in data["anomaly_findings"]:
        afp = af.get("properties") or {}
        cls = afp.get("anomaly_class", "?")
        A(f"### {cls}")
        A("")
        A(f"- **Severity:** {afp.get('aggregate_severity', '?')}")
        A(f"- **Confidence:** {afp.get('detection_confidence', '?')}")
        A(f"- **Primary bidders implicated:** {', '.join(afp.get('primary_bidder_names') or [])}")
        A(f"- **Cross-tender consistency:** {afp.get('cross_tender_consistency')} "
          f"({afp.get('cross_tender_appearances', 0)} of 3 tenders)")
        A("")
        A(f"**Decision reason:** {afp.get('decision_reason', '')}")
        A("")
        A("**Signal evidence:**")
        A("")
        A("| Signal type | Severity | Evidence | Citation |")
        A("|---|---|---|---|")
        for s in afp.get("signals") or []:
            A(f"| {s.get('signal_type', '?')} | {s.get('severity', '?')} | "
              f"{(s.get('evidence') or '').replace('|', '/')[:200]} | "
              f"_{(s.get('citation_source') or '').replace('|', '/')[:120]}_ |")
        A("")
        A(f"**Recommendation:** {afp.get('recommendation', '')}")
        A("")
        A(f"_Drilldown — BidAnomalyFinding node: `{af.get('node_id')}`_")
        A("")

    # === PART F — Committee Recommendation ===
    A("## Part F — Committee Recommendation")
    A("")
    eff = data["effective_l1"]
    if eff["entry"] is None:
        A("**⚠ NO EFFECTIVE L1**: every QUALIFIED bidder was rejected via "
          "ALB-rejection or cartel-suspect flagging. Recommend re-tender.")
    else:
        e = eff["entry"]
        A(f"**Effective L1 (post-anomaly adjustment):** **{e['bidder_name']}** "
          f"at **₹{e['bid_amount_cr']:.2f}cr** ({e['premium_pct']:.2f}% premium)")
        A("")
        A(f"_Rationale: {eff['rationale']}_")
    A("")
    A("### Three recommended options (committee to choose):")
    A("")
    for opt in data["recommendation_options"]:
        marker = " ← **RECOMMENDED**" if opt.get("preferred") else ""
        A(f"- **Option ({opt['option']}):** {opt['label']}{marker}")
    A("")
    A("**Committee Decision:**")
    A("")
    A("> ___________________________________________________________________________")
    A(">")
    A("> ___________________________________________________________________________")
    A(">")
    A("> Authorised Signatory (Committee Chair): _________________  Date: ____________")
    A("")

    # === PART G — Audit Trail ===
    A("## Part G — Audit Trail")
    A("")
    audit = data["audit"]
    A(f"- **Audit ID:** `{data['audit_id']}` "
      f"(SHA256 of sorted finding node_ids; deterministic across re-runs of identical inputs)")
    A(f"- **Findings consumed (total):** {audit['findings_consumed_count']}")
    A(f"  - BidEvaluationFinding rows: {audit['findings_consumed_breakdown']['BidEvaluationFinding']}")
    A(f"  - EligibilityMatrix rows: {audit['findings_consumed_breakdown']['EligibilityMatrix']}")
    A(f"  - TenderRanking rows: {audit['findings_consumed_breakdown']['TenderRanking']}")
    A(f"  - BidAnomalyFinding rows: {audit['findings_consumed_breakdown']['BidAnomalyFinding']}")
    A(f"- **Rules cited in evaluation:** {', '.join(f'`{r}`' for r in audit['rules_cited'])}")
    A("")
    A(f"- **TenderRanking node_id (drilldown):** `{data['drilldown']['tender_ranking_node_id']}`")
    A(f"- **EligibilityMatrix node_ids (8):** {', '.join(f'`{x}`' for x in data['drilldown']['eligibility_matrix_node_ids'])}")
    A(f"- **BidAnomalyFinding node_ids (2):** {', '.join(f'`{x}`' for x in data['drilldown']['anomaly_finding_node_ids'])}")
    A("")
    A(f"**Report template version:** `{REPORT_TEMPLATE_VERSION}`  ")
    A(f"**Generator:** `{SOURCE_REF}`  ")
    A("")
    A("_All claims in this report trace to a kg_node UUID. The 5-layer "
      "drilldown chain (ComparativeStatement → BidAnomalyFinding / "
      "TenderRanking / EligibilityMatrix → BidEvaluationFinding → "
      "BidSubmission / BidderProfile → fact_sheets) enables full citation "
      "verification by querying the listed node_ids directly._")
    A("")

    return "\n".join(md)


# ── DOCX rendering (python-docx) ─────────────────────────────────────

def render_docx(data: dict, out_path: Path) -> None:
    """Render the report as a DOCX. python-docx 1.2.0 is already installed."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()
    # Set default style sizes a touch smaller for table fit
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    t = data["tender"]

    # === Header block ===
    title = doc.add_heading("Comparative Statement of Bids", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run(f"Tender: {t['tender_name']}").bold = True
    doc.add_paragraph(f"NIT No.: {t['tender_nit_no']}")
    doc.add_paragraph(f"Estimated Contract Value: ₹{t['tender_ecv_cr']:.2f} crore")
    doc.add_paragraph(f"Tender Method: {t['tender_method']} (lowest bid wins among QUALIFIED)")
    doc.add_paragraph(f"Report generated: {data['report_generated_at']}")
    p = doc.add_paragraph()
    p.add_run("Audit ID: ").bold = True
    p.add_run(data["audit_id"]).font.name = "Courier New"

    # === PART A ===
    doc.add_heading("Part A — Tender Summary", level=1)
    doc.add_paragraph(f"Tender ID (internal): {t['tender_id']}", style="List Bullet")
    doc.add_paragraph(f"Estimated Contract Value (ECV): ₹{t['tender_ecv_cr']:.2f} crore",
                      style="List Bullet")
    doc.add_paragraph(f"Tender method: {t['tender_method']}", style="List Bullet")

    # === PART B ===
    doc.add_heading("Part B — Bidder Participation Overview", level=1)
    counts = data["participation_counts"]
    p = doc.add_paragraph()
    p.add_run(f"Total bidders received: {counts['total']}").bold = True
    doc.add_paragraph(f"QUALIFIED: {counts['qualified']}", style="List Bullet")
    doc.add_paragraph(f"FLAGGED for committee review: {counts['flagged']}", style="List Bullet")
    doc.add_paragraph(f"MARKED for documentation review: {counts['mark_for_doc']}", style="List Bullet")
    doc.add_paragraph(f"DISQUALIFIED: {counts['disqualified']}", style="List Bullet")

    doc.add_heading("Excluded bidders", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Bidder", "Aggregate verdict", "Exclusion summary",
                           "EligibilityMatrix node_id"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for ex in data["excluded_summaries"]:
        row = tbl.add_row().cells
        row[0].text = ex["bidder_name"]
        row[1].text = AGGREGATE_LABEL.get(ex["aggregate_verdict"], ex["aggregate_verdict"])
        row[2].text = ex["exclusion_summary"]
        row[3].text = ex["eligibility_matrix_node_id"]

    # === PART C ===
    doc.add_heading("Part C — Per-Bidder Detailed Evaluation", level=1)
    for bidder_data in data["per_bidder_evaluations"]:
        prof = bidder_data["profile"]
        em = bidder_data["em_props"]
        doc.add_heading(prof.get("company_name", prof.get("profile_id")), level=2)
        doc.add_paragraph(f"Contractor class: {prof.get('contractor_class', '?')}", style="List Bullet")
        doc.add_paragraph(f"Registration: {prof.get('registration_certificate_no', '?')} "
                          f"(state: {prof.get('registration_state', '?')})", style="List Bullet")
        doc.add_paragraph(f"PAN / GSTIN: {prof.get('pan', '?')} / {prof.get('gstin', '?')}",
                          style="List Bullet")
        doc.add_paragraph(f"Communication address: {prof.get('communication_address', '?')}",
                          style="List Bullet")
        doc.add_paragraph(f"Authorized signatory: {prof.get('authorized_signatory_name', '?')} "
                          f"({prof.get('authorized_signatory_role', '?')})", style="List Bullet")
        p = doc.add_paragraph()
        run = p.add_run(f"Aggregate verdict: "
                        f"{AGGREGATE_LABEL.get(em.get('aggregate_verdict'), em.get('aggregate_verdict'))} "
                        f"({em.get('count_qualified', 0)} QUALIFIED + "
                        f"{em.get('count_ineligible_hard_block', 0)} HARD_BLOCK + "
                        f"{em.get('count_ineligible_warning', 0)} WARNING + "
                        f"{em.get('count_gap', 0)} GAP)")
        run.bold = True
        # italicised reasoning
        p = doc.add_paragraph()
        p.add_run(em.get("aggregate_reasoning", "")).italic = True

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["#", "Criterion", "Verdict",
                               "Decision reason", "Rule", "Finding node_id"]):
            hdr[i].paragraphs[0].add_run(h).bold = True
        for i, f in enumerate(bidder_data["findings"], 1):
            fp = f.get("properties") or {}
            v = fp.get("verdict") or "?"
            cons = fp.get("evaluation_consequence") or ""
            row = tbl.add_row().cells
            row[0].text = str(i)
            row[1].text = fp.get("typology_code", "?")
            row[2].text = VERDICT_LABEL.get(v, v)
            reason = (fp.get("decision_reason") or "")[:200]
            row[3].text = reason
            row[4].text = fp.get("rule_id", "?")
            row[5].text = f.get("node_id", "?")
            # Bold HARD_BLOCK INELIGIBLE rows
            if cons == "HARD_BLOCK" and v == "INELIGIBLE":
                for cell in row:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    # === PART D ===
    doc.add_heading("Part D — Ranking of QUALIFIED Bidders", level=1)
    tr_props = data["tender_ranking_props"]
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Rank", "Bidder", "Bid (₹cr)", "Premium % vs ECV",
                           "ALB flag", "Distance from L1"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for r in tr_props.get("ranking") or []:
        row = tbl.add_row().cells
        row[0].text = r["rank_position"]
        row[1].text = r["bidder_name"]
        row[2].text = f"{r['bid_amount_cr']:.2f}"
        row[3].text = f"{r['premium_pct']:.2f}"
        row[4].text = "⚠ YES" if r.get("alb_flag") else "—"
        d_cr = r.get("distance_from_l1_cr", 0.0) or 0.0
        d_pct = r.get("distance_from_l1_pct", 0.0) or 0.0
        row[5].text = "—" if r["rank_position"] == "L1" else f"+₹{d_cr:.2f}cr ({d_pct:.2f}%)"
        # Bold L1
        if r["rank_position"] == "L1":
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_heading("ALB (Abnormally Low Bid) Detection", level=2)
    doc.add_paragraph(f"Methodology: {tr_props.get('alb_threshold_method', '?')}", style="List Bullet")
    doc.add_paragraph(f"Average of qualified bids: ₹{tr_props.get('average_qualified_bid_cr', 0.0):.2f}cr",
                      style="List Bullet")
    doc.add_paragraph(f"ALB threshold: ₹{tr_props.get('alb_threshold_cr', 0.0):.2f}cr "
                      f"(average × {tr_props.get('alb_multiplier', 0.80)})", style="List Bullet")
    doc.add_paragraph(f"ALB candidates: {tr_props.get('alb_candidates') or '(none)'}", style="List Bullet")
    p = doc.add_paragraph(style="List Bullet")
    if tr_props.get("alb_action_required"):
        run = p.add_run("Action required on L1: YES — L1 is ALB candidate")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    else:
        p.add_run("Action required on L1: No")
    if tr_props.get("alb_methodology_note"):
        p = doc.add_paragraph()
        p.add_run(f"Methodology note: {tr_props['alb_methodology_note'][:500]}").italic = True
    if tr_props.get("l1_l2_gap_cr") is not None:
        p = doc.add_paragraph()
        p.add_run(f"L1 → L2 gap: ₹{tr_props['l1_l2_gap_cr']:.2f}cr "
                  f"({tr_props.get('l1_l2_gap_pct', 0.0):.2f}%)").bold = True

    # === PART E ===
    doc.add_heading("Part E — Anomaly Findings", level=1)
    if not data["anomaly_findings"]:
        doc.add_paragraph("No anomaly findings emitted by CrossBidAnomalyDetector for this tender.")
    for af in data["anomaly_findings"]:
        afp = af.get("properties") or {}
        cls = afp.get("anomaly_class", "?")
        doc.add_heading(cls, level=2)
        doc.add_paragraph(f"Severity: {afp.get('aggregate_severity', '?')}", style="List Bullet")
        doc.add_paragraph(f"Confidence: {afp.get('detection_confidence', '?')}", style="List Bullet")
        doc.add_paragraph(f"Primary bidders implicated: "
                          f"{', '.join(afp.get('primary_bidder_names') or [])}",
                          style="List Bullet")
        doc.add_paragraph(f"Cross-tender consistency: {afp.get('cross_tender_consistency')} "
                          f"({afp.get('cross_tender_appearances', 0)} of 3 tenders)",
                          style="List Bullet")
        p = doc.add_paragraph()
        p.add_run("Decision reason: ").bold = True
        p.add_run(afp.get("decision_reason", ""))
        p = doc.add_paragraph()
        p.add_run("Signal evidence:").bold = True
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Signal type", "Severity", "Evidence", "Citation"]):
            hdr[i].paragraphs[0].add_run(h).bold = True
        for s in afp.get("signals") or []:
            row = tbl.add_row().cells
            row[0].text = s.get("signal_type", "?")
            row[1].text = s.get("severity", "?")
            row[2].text = (s.get("evidence") or "")[:300]
            cite_para = row[3].paragraphs[0]
            cite_para.add_run((s.get("citation_source") or "")[:200]).italic = True
        p = doc.add_paragraph()
        p.add_run("Recommendation: ").bold = True
        p.add_run(afp.get("recommendation", ""))
        p = doc.add_paragraph()
        p.add_run(f"Drilldown — BidAnomalyFinding node: {af.get('node_id')}").italic = True

    # === PART F ===
    doc.add_heading("Part F — Committee Recommendation", level=1)
    eff = data["effective_l1"]
    p = doc.add_paragraph()
    if eff["entry"] is None:
        run = p.add_run("⚠ NO EFFECTIVE L1: every QUALIFIED bidder was rejected via "
                        "ALB-rejection or cartel-suspect flagging. Recommend re-tender.")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    else:
        e = eff["entry"]
        run = p.add_run(f"Effective L1 (post-anomaly adjustment): {e['bidder_name']} "
                        f"at ₹{e['bid_amount_cr']:.2f}cr ({e['premium_pct']:.2f}% premium)")
        run.bold = True
    p = doc.add_paragraph()
    p.add_run(f"Rationale: {eff['rationale']}").italic = True

    doc.add_heading("Three recommended options (committee to choose):", level=2)
    for opt in data["recommendation_options"]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"Option ({opt['option']}): {opt['label']}")
        if opt.get("preferred"):
            run.bold = True
            p.add_run(" ← RECOMMENDED").bold = True

    doc.add_heading("Committee Decision", level=2)
    doc.add_paragraph("___________________________________________________________________________")
    doc.add_paragraph("___________________________________________________________________________")
    doc.add_paragraph("")
    doc.add_paragraph("Authorised Signatory (Committee Chair): _________________  Date: ____________")

    # === PART G ===
    doc.add_heading("Part G — Audit Trail", level=1)
    audit = data["audit"]
    p = doc.add_paragraph()
    p.add_run(f"Audit ID: {data['audit_id']}").bold = True
    doc.add_paragraph("(SHA256 of sorted finding node_ids; deterministic across re-runs of identical inputs)")
    p = doc.add_paragraph()
    p.add_run(f"Findings consumed (total): {audit['findings_consumed_count']}").bold = True
    fb = audit["findings_consumed_breakdown"]
    doc.add_paragraph(f"BidEvaluationFinding rows: {fb['BidEvaluationFinding']}", style="List Bullet")
    doc.add_paragraph(f"EligibilityMatrix rows: {fb['EligibilityMatrix']}", style="List Bullet")
    doc.add_paragraph(f"TenderRanking rows: {fb['TenderRanking']}", style="List Bullet")
    doc.add_paragraph(f"BidAnomalyFinding rows: {fb['BidAnomalyFinding']}", style="List Bullet")
    p = doc.add_paragraph()
    p.add_run(f"Rules cited: {', '.join(audit['rules_cited'])}").italic = True

    p = doc.add_paragraph()
    p.add_run(f"TenderRanking node_id: {data['drilldown']['tender_ranking_node_id']}").italic = True
    p = doc.add_paragraph()
    p.add_run(f"EligibilityMatrix node_ids (8): "
              f"{', '.join(data['drilldown']['eligibility_matrix_node_ids'])}").italic = True
    p = doc.add_paragraph()
    p.add_run(f"BidAnomalyFinding node_ids (2): "
              f"{', '.join(data['drilldown']['anomaly_finding_node_ids'])}").italic = True

    p = doc.add_paragraph()
    p.add_run(f"Report template version: {REPORT_TEMPLATE_VERSION}\n"
              f"Generator: {SOURCE_REF}").italic = True

    p = doc.add_paragraph()
    p.add_run("All claims in this report trace to a kg_node UUID. The 5-layer "
              "drilldown chain (ComparativeStatement → BidAnomalyFinding / "
              "TenderRanking / EligibilityMatrix → BidEvaluationFinding → "
              "BidSubmission / BidderProfile → fact_sheets) enables full "
              "citation verification by querying the listed node_ids "
              "directly.").italic = True

    doc.save(out_path)


# ── Idempotent cleanup ────────────────────────────────────────────────

def _delete_prior_comparative_statements() -> int:
    rows = rest_get("kg_nodes", {
        "select":     "node_id",
        "node_type":  f"eq.{TYPOLOGY}",
        "source_ref": f"eq.{SOURCE_REF}",
    })
    for row in rows:
        rest_delete("kg_nodes", {"node_id": f"eq.{row['node_id']}"})
    return len(rows)


# ── Main ──────────────────────────────────────────────────────────────

def main() -> int:
    t_start = time.perf_counter()
    print("=" * 76)
    print(f"  Sub-block 7 — ComparativeStatementGenerator (visible demo artifact)")
    print(f"  source_ref : {SOURCE_REF}")
    print(f"  artifact_dir: {ARTIFACT_DIR}")
    print("=" * 76)

    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)

    n_prior = _delete_prior_comparative_statements()
    if n_prior:
        print(f"  cleanup: removed {n_prior} prior ComparativeStatement row(s)")

    sentinel_pre = {
        "ValidationFinding":    rest_count("kg_nodes", {"node_type": "eq.ValidationFinding"}),
        "BidEvaluationFinding": rest_count("kg_nodes", {"node_type": "eq.BidEvaluationFinding"}),
        "BIDDER_VIOLATES_RULE": rest_count("kg_edges", {"edge_type": "eq.BIDDER_VIOLATES_RULE"}),
        "EligibilityMatrix":    rest_count("kg_nodes", {"node_type": "eq.EligibilityMatrix"}),
        "TenderRanking":        rest_count("kg_nodes", {"node_type": "eq.TenderRanking"}),
        "BidAnomalyFinding":    rest_count("kg_nodes", {"node_type": "eq.BidAnomalyFinding"}),
    }
    print(f"\n── Sentinel snapshot (pre) ──")
    for k, v in sentinel_pre.items():
        print(f"  {k:24s} : {v}")

    # ── Load upstream sources once ──
    print(f"\n── Load upstream sources ──")
    tr_rows  = fetch_all_by_type("TenderRanking")
    em_rows  = fetch_all_by_type("EligibilityMatrix")
    af_rows  = fetch_all_by_type("BidAnomalyFinding")
    bp_rows  = fetch_all_by_type("BidderProfile")
    print(f"  TenderRanking      : {len(tr_rows)}")
    print(f"  EligibilityMatrix  : {len(em_rows)}")
    print(f"  BidAnomalyFinding  : {len(af_rows)}")
    print(f"  BidderProfile      : {len(bp_rows)}")

    bp_by_id = {bp["doc_id"]: bp for bp in bp_rows}
    em_by_bid = {em["properties"].get("bid_submission_id"): em for em in em_rows}

    # Bulk fetch all BidEvaluationFinding rows once (240)
    all_finding_ids: list[str] = []
    for em in em_rows:
        all_finding_ids.extend(em["properties"].get("finding_node_ids") or [])
    print(f"  BidEvaluationFinding (via EM drilldown): fetching {len(all_finding_ids)} UUIDs...")
    bef_rows = fetch_findings_by_node_ids(all_finding_ids)
    bef_by_id = {f["node_id"]: f for f in bef_rows}
    print(f"  BidEvaluationFinding : {len(bef_rows)} fetched")

    # ── Per-tender report generation ──
    emitted: list[dict] = []
    print()
    for tr in sorted(tr_rows, key=lambda x: (x["properties"].get("tender_id") or "")):
        trp = tr["properties"] or {}
        tid = trp.get("tender_id")
        print(f"── Tender: {tid} ──")

        # Filter sources to this tender
        ems_this = [e for e in em_rows
                    if (e["properties"] or {}).get("tender_id") == tid]
        afs_this = [a for a in af_rows
                    if (a["properties"] or {}).get("tender_id") == tid]
        print(f"  EligibilityMatrix this tender: {len(ems_this)}")
        print(f"  Anomalies this tender        : {len(afs_this)}")

        # Effective L1
        eff_entry, alb_rej, cartel_ref = compute_effective_l1(trp, afs_this)
        if eff_entry:
            rationale = []
            if alb_rej:
                rationale.append(f"L1 ({trp.get('l1_winner_bidder_name')}) ALB-rejected per CVC norms")
            if cartel_ref:
                rationale.append(f"cartel-suspect bidders ({', '.join(cartel_ref)}) referred to committee")
            rationale.append(f"first non-skipped bidder in ranking = {eff_entry['bidder_name']}")
            rationale_str = "; ".join(rationale) + "."
        else:
            rationale_str = ("No effective L1: all QUALIFIED bidders skipped "
                             "(ALB-rejection + cartel-referral exhausted ranking).")

        # Recommendation options
        recommendation_options = []
        if alb_rej:
            recommendation_options.append({
                "option": "i",
                "label": (f"Reject ALB-flagged L1 ({trp.get('l1_winner_bidder_name')}) "
                          f"per CVC ALB norms; proceed with "
                          f"{eff_entry['bidder_name'] if eff_entry else 'no remaining bidder'} as effective L1"),
                "preferred": True if eff_entry else False,
            })
            recommendation_options.append({
                "option": "ii",
                "label": (f"Require cost-build-up + bank-guarantee justification "
                          f"from {trp.get('l1_winner_bidder_name')} per CVC ALB norms; "
                          f"re-decide after review"),
                "preferred": False,
            })
        recommendation_options.append({
            "option": "iii",
            "label": "Refer all bids for re-evaluation (or re-tender if no L1 emerges)",
            "preferred": False,
        })

        # Per-bidder evaluations — ordered by company_name for consistency
        per_bidder = []
        for em in sorted(ems_this,
                         key=lambda x: x["properties"].get("bidder_profile_id") or ""):
            emp = em["properties"] or {}
            bpid = emp.get("bidder_profile_id")
            bp = bp_by_id.get(bpid)
            if bp is None:
                continue
            findings = []
            for fid in emp.get("finding_node_ids") or []:
                f = bef_by_id.get(fid)
                if f:
                    findings.append(f)
            # Sort by typology_code for consistent ordering
            findings.sort(key=lambda f: (f.get("properties") or {}).get("typology_code") or "")
            per_bidder.append({
                "profile":  bp.get("properties") or {},
                "em_props": emp,
                "findings": findings,
            })

        # Excluded-bidder summary (from non-QUALIFIED EligibilityMatrix rows)
        excluded_summaries = []
        for em in ems_this:
            emp = em["properties"] or {}
            v = emp.get("aggregate_verdict")
            if v == "QUALIFIED":
                continue
            n_hb = emp.get("count_ineligible_hard_block", 0)
            n_w  = emp.get("count_ineligible_warning", 0)
            n_g  = emp.get("count_gap", 0)
            if v == "DISQUALIFIED":
                ex_sum = f"DISQUALIFIED — {n_hb} HARD_BLOCK failure(s)"
                if n_w:
                    ex_sum += f" + {n_w} WARNING"
            elif v == "FLAGGED_FOR_COMMITTEE_REVIEW":
                ex_sum = f"FLAGGED — {n_w} WARNING finding(s) require committee review"
            elif v == "MARK_FOR_DOCUMENTATION_REVIEW":
                ex_sum = f"MARK_FOR_DOC — {n_g} GAP finding(s) require bidder documentation"
            else:
                ex_sum = v or "EXCLUDED"
            excluded_summaries.append({
                "bidder_name":              emp.get("bidder_name"),
                "bidder_profile_id":        emp.get("bidder_profile_id"),
                "aggregate_verdict":        v,
                "exclusion_summary":        ex_sum,
                "eligibility_matrix_node_id": em["node_id"],
            })
        # Participation counts
        verdict_counts = defaultdict(int)
        for em in ems_this:
            verdict_counts[em["properties"].get("aggregate_verdict")] += 1
        participation_counts = {
            "total":         len(ems_this),
            "qualified":     verdict_counts.get("QUALIFIED", 0),
            "flagged":       verdict_counts.get("FLAGGED_FOR_COMMITTEE_REVIEW", 0),
            "mark_for_doc":  verdict_counts.get("MARK_FOR_DOCUMENTATION_REVIEW", 0),
            "disqualified":  verdict_counts.get("DISQUALIFIED", 0),
        }

        # Audit chain — collect all finding_node_ids consumed for audit_id
        finding_ids_for_audit: list[str] = []
        for em in ems_this:
            finding_ids_for_audit.extend(em["properties"].get("finding_node_ids") or [])
            finding_ids_for_audit.append(em["node_id"])
        finding_ids_for_audit.append(tr["node_id"])
        for af in afs_this:
            finding_ids_for_audit.append(af["node_id"])

        # Rules cited — union across all this tender's BidEvaluationFinding rule_ids
        rules_set: set[str] = set()
        for bd in per_bidder:
            for f in bd["findings"]:
                rid = (f.get("properties") or {}).get("rule_id")
                if rid:
                    rules_set.add(rid)
                srid = (f.get("properties") or {}).get("secondary_rule_id")
                if srid:
                    rules_set.add(srid)
        # Also from anomaly findings (citation_source is text, not rule_id; skip)
        rules_cited = sorted(rules_set)

        # Build report data dict
        report_data = {
            "tender": {
                "tender_id":       tid,
                "tender_name":     trp.get("tender_name"),
                "tender_ecv_cr":   trp.get("tender_ecv_cr"),
                "tender_nit_no":   trp.get("tender_nit_no"),
                "tender_method":   trp.get("tender_method"),
            },
            "report_generated_at":   _dt.datetime.now(tz=_dt.timezone.utc).isoformat(timespec="seconds"),
            "participation_counts":  participation_counts,
            "excluded_summaries":    excluded_summaries,
            "per_bidder_evaluations": per_bidder,
            "tender_ranking_props":  trp,
            "anomaly_findings":      afs_this,
            "effective_l1": {
                "entry":                eff_entry,
                "rationale":            rationale_str,
                "alb_rejected_ids":     alb_rej,
                "cartel_referred_ids":  cartel_ref,
            },
            "recommendation_options": recommendation_options,
            "audit": {
                "findings_consumed_count": (len([f for bd in per_bidder for f in bd["findings"]])
                                            + len(ems_this) + 1 + len(afs_this)),
                "findings_consumed_breakdown": {
                    "BidEvaluationFinding": sum(len(bd["findings"]) for bd in per_bidder),
                    "EligibilityMatrix":    len(ems_this),
                    "TenderRanking":         1,
                    "BidAnomalyFinding":     len(afs_this),
                },
                "rules_cited": rules_cited,
            },
            "drilldown": {
                "tender_ranking_node_id":     tr["node_id"],
                "eligibility_matrix_node_ids": [e["node_id"] for e in ems_this],
                "anomaly_finding_node_ids":    [a["node_id"] for a in afs_this],
            },
        }
        # audit_id computed from finding_ids_for_audit
        report_data["audit_id"] = compute_audit_id(finding_ids_for_audit)

        # Render artifacts
        md_text  = render_markdown(report_data)
        md_path  = ARTIFACT_DIR / f"{tid}.md"
        docx_path = ARTIFACT_DIR / f"{tid}.docx"
        md_path.write_text(md_text, encoding="utf-8")
        render_docx(report_data, docx_path)
        print(f"  → md   : {md_path}  ({md_path.stat().st_size} bytes)")
        print(f"  → docx : {docx_path}  ({docx_path.stat().st_size} bytes)")

        # Emit ComparativeStatement kg_node
        eff_id = eff_entry["bidder_profile_id"] if eff_entry else None
        eff_name = eff_entry["bidder_name"] if eff_entry else None
        eff_amount = eff_entry["bid_amount_cr"] if eff_entry else None

        cartel_pairs_summary = []
        alb_corroboration_summary = []
        for af in afs_this:
            afp = af.get("properties") or {}
            cls = afp.get("anomaly_class")
            if cls == "CARTEL_SUSPECT":
                cartel_pairs_summary.append({
                    "bidder_ids":               afp.get("primary_bidder_ids") or [],
                    "signal_count":             afp.get("signal_count"),
                    "severity":                 afp.get("aggregate_severity"),
                    "anomaly_finding_node_id":  af["node_id"],
                })
            elif cls == "ALB_CORROBORATION":
                alb_corroboration_summary.append({
                    "bidder_id":                  (afp.get("primary_bidder_ids") or [None])[0],
                    "cross_tender_appearances":   afp.get("cross_tender_appearances"),
                    "anomaly_finding_node_id":    af["node_id"],
                })

        cs_props = {
            "tier":                  7,
            "tender_id":             tid,
            "tender_name":           report_data["tender"]["tender_name"],
            "tender_ecv_cr":         report_data["tender"]["tender_ecv_cr"],
            "tender_nit_no":         report_data["tender"]["tender_nit_no"],
            "tender_method":         report_data["tender"]["tender_method"],
            "report_generated_at":   report_data["report_generated_at"],
            "audit_id":              report_data["audit_id"],

            # Participation
            "total_bidders":         participation_counts["total"],
            "qualified_count":       participation_counts["qualified"],
            "flagged_count":         participation_counts["flagged"],
            "mark_for_doc_count":    participation_counts["mark_for_doc"],
            "disqualified_count":    participation_counts["disqualified"],

            # Raw L1
            "l1_winner_bidder_id":   trp.get("l1_winner_bidder_id"),
            "l1_winner_bidder_name": trp.get("l1_winner_bidder_name"),
            "l1_amount_cr":          trp.get("l1_amount_cr"),
            "l1_alb_flag":           bool(trp.get("alb_action_required")),

            # Anomalies
            "cartel_suspect_pairs":  cartel_pairs_summary,
            "alb_corroboration_bidders": alb_corroboration_summary,

            # Effective L1
            "effective_l1_bidder_id":   eff_id,
            "effective_l1_bidder_name": eff_name,
            "effective_l1_amount_cr":   eff_amount,
            "effective_l1_rationale":   rationale_str,
            "alb_rejected_bidder_ids":  alb_rej,
            "cartel_referred_bidder_ids": cartel_ref,

            # Committee recommendation
            "committee_recommendation_summary": (
                f"Recommend {recommendation_options[0]['label']}."
                if recommendation_options else
                "No automated recommendation; manual review required."
            ),
            "recommendation_options": recommendation_options,

            # Audit / drilldown
            "findings_consumed_count":  report_data["audit"]["findings_consumed_count"],
            "findings_consumed_breakdown": report_data["audit"]["findings_consumed_breakdown"],
            "rules_cited":              rules_cited,
            "tender_ranking_node_id":   tr["node_id"],
            "eligibility_matrix_node_ids": [e["node_id"] for e in ems_this],
            "anomaly_finding_node_ids":    [a["node_id"] for a in afs_this],

            # Artifacts (Path A — Markdown + DOCX; PDF deferred to L75)
            "md_artifact_path":      str(md_path),
            "docx_artifact_path":    str(docx_path),
            "pdf_artifact_path":     None,
            "pdf_artifact_status":   "deferred_no_renderer_in_env",
            "pdf_followup_options":  ["reportlab", "weasyprint"],

            # Metadata
            "extracted_by":            SOURCE_REF,
            "report_template_version": REPORT_TEMPLATE_VERSION,
        }

        if eff_entry is None:
            label = (f"ComparativeStatement: {tid.replace('tender_synth_', '')} "
                     f"— NO EFFECTIVE L1 (all bidders skipped)")
        else:
            label = (f"ComparativeStatement: {tid.replace('tender_synth_', '')} "
                     f"— Effective L1 {eff_name} @ ₹{eff_amount:.2f}cr "
                     f"(raw L1 {trp.get('l1_winner_bidder_name')} ALB-rejected)")

        inserted = rest_post("kg_nodes", [{
            "doc_id":     tid,
            "node_type":  TYPOLOGY,
            "label":      label,
            "properties": cs_props,
            "source_ref": SOURCE_REF,
        }])[0]
        emitted.append(inserted)
        print(f"  → ComparativeStatement {inserted['node_id']}")
        print(f"    effective_L1: {eff_name} @ ₹{eff_amount}cr" if eff_entry
              else "    effective_L1: NONE")
        print(f"    audit_id: {report_data['audit_id']}")
        print()

    # ── Sentinel post-snapshot ──
    sentinel_post = {
        "ValidationFinding":    rest_count("kg_nodes", {"node_type": "eq.ValidationFinding"}),
        "BidEvaluationFinding": rest_count("kg_nodes", {"node_type": "eq.BidEvaluationFinding"}),
        "BIDDER_VIOLATES_RULE": rest_count("kg_edges", {"edge_type": "eq.BIDDER_VIOLATES_RULE"}),
        "EligibilityMatrix":    rest_count("kg_nodes", {"node_type": "eq.EligibilityMatrix"}),
        "TenderRanking":        rest_count("kg_nodes", {"node_type": "eq.TenderRanking"}),
        "BidAnomalyFinding":    rest_count("kg_nodes", {"node_type": "eq.BidAnomalyFinding"}),
        "ComparativeStatement": rest_count("kg_nodes", {"node_type": f"eq.{TYPOLOGY}"}),
    }
    print(f"── Sentinel snapshot (post) ──")
    drift = False
    for k, v in sentinel_post.items():
        pre_v = sentinel_pre.get(k)
        marker = ""
        if pre_v is not None and pre_v != v:
            marker = f"  ⚠ DRIFT (was {pre_v})"
            drift = True
        print(f"  {k:24s} : {v}{marker}")
    if drift:
        print(f"  ✗ sentinel drift — upstream tables modified during generator run")
        return 2

    wall = time.perf_counter() - t_start
    print()
    print("=" * 76)
    print(f"  ComparativeStatementGenerator complete — emitted {len(emitted)} row(s) "
          f"+ {len(emitted) * 2} artifact file(s) in {wall:.2f}s")
    print("=" * 76)
    return 0


if __name__ == "__main__":
    from modules.validation.verdict_emitter import main_with_crash_resilience
    raise SystemExit(main_with_crash_resilience(
        main, doc_id=AGGREGATOR_DOC_ID, typology=TYPOLOGY))
