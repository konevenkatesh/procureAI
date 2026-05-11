"""Shared helpers for M4 drafters (M4.2)."""
from __future__ import annotations

import hashlib
import json
import sys
import time
from pathlib import Path
from typing import Iterable

import requests

REPO = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(REPO))

from builder.config import settings  # noqa: E402

REST = settings.supabase_rest_url
H = {"apikey": settings.supabase_anon_key,
     "Authorization": f"Bearer {settings.supabase_anon_key}"}

ARTIFACT_DIR = Path("/tmp/m4_drafts")
ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)


# ── REST helpers (with retry) ─────────────────────────────────────────

def _request_with_retry(method, url, **kw):
    last_exc = None
    for i in range(4):
        try:
            r = requests.request(method, url, timeout=30, **kw)
            r.raise_for_status()
            return r
        except (requests.exceptions.ConnectionError,
                requests.exceptions.Timeout) as exc:
            last_exc = exc
            if i == 3:
                break
            time.sleep(0.5 * (2 ** i))
    raise last_exc  # type: ignore[misc]


def rest_get(path, params=None):
    return _request_with_retry("GET", f"{REST}/rest/v1/{path}",
                               params=params or {}, headers=H).json()


def rest_get_range(path, params=None, range_header="0-1000"):
    return _request_with_retry("GET", f"{REST}/rest/v1/{path}",
                               params=params or {},
                               headers={**H, "Range": range_header}).json()


def rest_post(path, body):
    r = _request_with_retry("POST", f"{REST}/rest/v1/{path}", json=body,
                            headers={**H, "Content-Type": "application/json",
                                     "Prefer": "return=representation"})
    return r.json()


def rest_delete(path, params=None):
    requests.delete(f"{REST}/rest/v1/{path}", params=params or {}, headers=H, timeout=15)


def rest_count(path, params=None):
    r = requests.get(f"{REST}/rest/v1/{path}", params=params or {},
                     headers={**H, "Prefer": "count=exact", "Range": "0-0"}, timeout=30)
    return int(r.headers["Content-Range"].split("/")[1])


# ── Audit ID computation (deterministic SHA256) ───────────────────────

def compute_audit_id(communication_type: str,
                     recipient_id: str | None,
                     tender_id: str,
                     source_finding_node_ids: Iterable[str]) -> str:
    """Deterministic 16-char hex SHA256 of the inputs. Identical inputs
    produce identical audit_ids across re-runs — enables idempotent
    re-emission."""
    payload = (f"{communication_type}|"
               f"{recipient_id or 'INTERNAL'}|"
               f"{tender_id}|"
               f"{','.join(sorted(source_finding_node_ids))}")
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]


# ── Lookups ──────────────────────────────────────────────────────────

_PROFILE_CACHE: dict[str, dict] = {}


def get_bidder_profile(profile_id: str) -> dict:
    """Cached BidderProfile lookup."""
    if profile_id in _PROFILE_CACHE:
        return _PROFILE_CACHE[profile_id]
    rows = rest_get_range("kg_nodes", {
        "select": "node_id,doc_id,properties",
        "doc_id": f"eq.{profile_id}",
        "node_type": "eq.BidderProfile",
    })
    _PROFILE_CACHE[profile_id] = rows[0] if rows else {}
    return _PROFILE_CACHE[profile_id]


_TENDER_INFO = {
    "tender_synth_kurnool": {"name": "District Hospital, Kurnool",
                              "nit_no": "100/PROC/APIIC/1/2026",
                              "ecv_cr": 85.0},
    "tender_synth_ja":      {"name": "Andhra Pradesh Judicial Academy",
                              "nit_no": "JA/2026/CW/001",
                              "ecv_cr": 125.5},
    "tender_synth_hc":      {"name": "Andhra Pradesh High Court complex",
                              "nit_no": "HC/APCRDA/2026/PROC/001",
                              "ecv_cr": 365.16},
}


def get_tender_info(tender_id: str) -> dict:
    return _TENDER_INFO.get(tender_id, {"name": tender_id, "nit_no": "?", "ecv_cr": 0.0})


# ── Idempotent cleanup ───────────────────────────────────────────────

def delete_prior_communications(communication_type: str, source_ref: str) -> int:
    """Delete all Communication kg_nodes of the given type emitted by the
    given source_ref. Idempotent — safe to call before re-emit."""
    rows = rest_get("kg_nodes", {
        "select": "node_id",
        "node_type": "eq.Communication",
        "properties->>communication_type": f"eq.{communication_type}",
        "source_ref": f"eq.{source_ref}",
    })
    for r in rows:
        rest_delete("kg_nodes", {"node_id": f"eq.{r['node_id']}"})
    return len(rows)


# ── Sentinel snapshot ────────────────────────────────────────────────

def snapshot_sentinels() -> dict[str, int]:
    return {
        "ValidationFinding":    rest_count("kg_nodes", {"node_type": "eq.ValidationFinding"}),
        "BidEvaluationFinding": rest_count("kg_nodes", {"node_type": "eq.BidEvaluationFinding"}),
        "EligibilityMatrix":    rest_count("kg_nodes", {"node_type": "eq.EligibilityMatrix"}),
        "TenderRanking":        rest_count("kg_nodes", {"node_type": "eq.TenderRanking"}),
        "BidAnomalyFinding":    rest_count("kg_nodes", {"node_type": "eq.BidAnomalyFinding"}),
        "ComparativeStatement": rest_count("kg_nodes", {"node_type": "eq.ComparativeStatement"}),
        "BIDDER_VIOLATES_RULE": rest_count("kg_edges", {"edge_type": "eq.BIDDER_VIOLATES_RULE"}),
        "Communication":        rest_count("kg_nodes", {"node_type": "eq.Communication"}),
    }


# ── DOCX rendering from Markdown (M4.3) ──────────────────────────────

def render_docx_from_md(content_md: str, out_path: Path, title: str) -> None:
    """Render Communication's Markdown body to a DOCX file.

    Simple structural conversion:
      # heading → Heading 1
      ## heading → Heading 2
      ### heading → Heading 3
      **bold** → bold run
      *italic* → italic run
      - / * bullet → List Bullet style
      `code` → Courier New run
      > quote → italic indented
      --- → page-width line (rendered as empty paragraph)
      otherwise → Normal paragraph
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title at top
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_runs(para, text):
        """Parse inline markdown (bold/italic/code) into runs."""
        # Pattern: matches **bold** or *italic* or `code`
        pat = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
        last_end = 0
        for m in pat.finditer(text):
            if m.start() > last_end:
                para.add_run(text[last_end:m.start()])
            token = m.group(0)
            if token.startswith("**") and token.endswith("**"):
                r = para.add_run(token[2:-2])
                r.bold = True
            elif token.startswith("`") and token.endswith("`"):
                r = para.add_run(token[1:-1])
                r.font.name = "Courier New"
                r.font.size = Pt(9)
            elif token.startswith("*") and token.endswith("*"):
                r = para.add_run(token[1:-1])
                r.italic = True
            else:
                para.add_run(token)
            last_end = m.end()
        if last_end < len(text):
            para.add_run(text[last_end:])

    for raw in content_md.split("\n"):
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            # Skip the # title at the top of the .md (already added as title)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.add_run("─" * 60).font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line[2:])
            continue
        if re.match(r"^\s*[-*]\s", line):
            content = re.sub(r"^\s*[-*]\s", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, content)
            continue
        if re.match(r"^\s*\d+\.\s", line):
            content = re.sub(r"^\s*\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            add_runs(p, content)
            continue
        # Default paragraph
        p = doc.add_paragraph()
        add_runs(p, line)

    doc.save(out_path)


def render_docx_for_communication(node_id: str, content_md: str,
                                   artifact_path_md: str,
                                   title: str) -> str:
    """Render DOCX next to the .md artifact and return DOCX path.

    Updates Communication kg_node properties.artifact_path_docx via
    fetch-modify-patch (Supabase JSONB needs full properties replacement
    to merge correctly — PostgREST doesn't natively merge nested JSONB).
    """
    md_path = Path(artifact_path_md)
    docx_path = md_path.with_suffix(".docx")
    render_docx_from_md(content_md, docx_path, title)
    # Fetch current properties
    r = requests.get(f"{REST}/rest/v1/kg_nodes",
                     params={"select": "properties", "node_id": f"eq.{node_id}"},
                     headers=H, timeout=30).json()
    if not r:
        return str(docx_path)
    props = r[0]["properties"] or {}
    props["artifact_path_docx"] = str(docx_path)
    # PATCH with full updated properties
    requests.patch(
        f"{REST}/rest/v1/kg_nodes",
        params={"node_id": f"eq.{node_id}"},
        headers={**H, "Content-Type": "application/json"},
        json={"properties": props},
        timeout=30,
    )
    return str(docx_path)


def assert_sentinel_preserved(pre: dict, post: dict, excluded_keys: tuple = ("Communication",)):
    """Verify all Module 3 sentinels unchanged. Communication count is
    expected to grow; everything else must hold."""
    drift = []
    for k in pre:
        if k in excluded_keys:
            continue
        if pre[k] != post[k]:
            drift.append(f"{k}: {pre[k]} → {post[k]}")
    if drift:
        raise RuntimeError(f"Sentinel drift detected: {'; '.join(drift)}")
    return True
