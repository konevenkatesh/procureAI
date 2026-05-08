"""Convert a draft tender Markdown file into a professionally typeset DOCX.

The Markdown produced by ``scripts/draft_tender.py`` uses a constrained
subset of GitHub-flavoured Markdown:

  - ATX headings (``#``..``######``)
  - Pipe tables (``| col | col |`` followed by ``|---|---|``)
  - Horizontal rules (``---``)
  - Blockquotes (``> ...``)  — used only for the cover/license panel
  - Inline emphasis: ``**bold**``, ``*italic*`` / ``_italic_``,
    backticks for ``code``
  - Inline links: ``[text](url)``
  - Plain paragraphs separated by blank lines
  - Bullet / numbered lists ("- item" or "1. item")

Beyond the markdown subset, this module adds:

  • A government-document **cover page** (seal placeholder, department
    name, NIT number, name-of-work, issue date, page break).
  • A **page footer** on every page — Contractor / page n / department
    + officer.
  • **Typography** matching real AP tenders — Times New Roman body,
    Arial Bold headings, justified text, NIT body table with
    proportional column widths.
  • A **structured forms renderer** (``render_form_section``) that
    detects rows inside Section IV and emits proper tender forms
    (Statement-I turnover, PBG proforma, LoA, declarations) instead of
    the raw 2-column markdown table the drafter ships.

Sections of this file:

  ── Part 1: parsing helpers + inline tokens
  ── Part 2: docx primitives (cell shading, borders, blanks)
  ── Part 3: header-param extraction (parse cover area)
  ── Part 4: cover page builder
  ── Part 5: page footer + typography
  ── Part 6: form classifier + render_form_section() + handlers
  ── Part 7: NIT body table styler
  ── Part 8: top-level md_to_docx() entry point
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Emu

# ──────────────────────────────────────────────────────────────────────
# Part 1 — Inline tokens + line-shape regexes
# ──────────────────────────────────────────────────────────────────────

_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*"          # **bold**
    r"|__[^_]+__"              # __bold__
    r"|\*[^*]+\*"              # *italic*
    r"|_[^_]+_"                # _italic_
    r"|`[^`]+`"                # `code`
    r"|\[[^\]]+\]\([^)]+\))"   # [text](url)
)
_LINK_RE       = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_HEADING_RE    = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")
_TABLE_SEP_RE  = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")
_BULLET_RE     = re.compile(r"^[\-\*\+]\s+(.*)$")
_NUMLIST_RE    = re.compile(r"^\d+\.\s+(.*)$")
_HR_RE         = re.compile(r"^\s*-{3,}\s*$")
# HTML comments — used by the fixed-skeleton sections to embed
# machine-readable markers like `<!-- {{BDS_REF: ITB_4.1}} -->` that
# the substitution layer reads but the rendered DOCX must not show.
# Single-line comments are skipped wholesale; inline comments are
# stripped from paragraph text before run-tokenisation.
_HTML_COMMENT_LINE_RE   = re.compile(r"^\s*<!--.*-->\s*$")
_HTML_COMMENT_INLINE_RE = re.compile(r"<!--.*?-->")


def _add_runs(paragraph, text: str) -> None:
    """Split text into Markdown inline tokens and append runs to paragraph."""
    if not text:
        return
    for chunk in _INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = paragraph.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith("__") and chunk.endswith("__"):
            r = paragraph.add_run(chunk[2:-2]); r.bold = True
        elif (chunk.startswith("*") and chunk.endswith("*")
              and len(chunk) >= 3):
            r = paragraph.add_run(chunk[1:-1]); r.italic = True
        elif (chunk.startswith("_") and chunk.endswith("_")
              and len(chunk) >= 3):
            r = paragraph.add_run(chunk[1:-1]); r.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            m = _LINK_RE.fullmatch(chunk)
            if m:
                r = paragraph.add_run(m.group(1))
                r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                r.underline = True
            else:
                paragraph.add_run(chunk)


def _strip_pipes(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"):   s = s[:-1]
    parts: list[str] = []
    cur, i = [], 0
    while i < len(s):
        ch = s[i]
        if ch == "\\" and i + 1 < len(s) and s[i+1] == "|":
            cur.append("|"); i += 2; continue
        if ch == "|":
            parts.append("".join(cur).strip())
            cur = []; i += 1; continue
        cur.append(ch); i += 1
    parts.append("".join(cur).strip())
    return parts


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


# ──────────────────────────────────────────────────────────────────────
# Part 2 — DOCX primitives
# ──────────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, *, sides=("top","left","bottom","right"),
                     sz: str = "4", color: str = "808080") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in sides:
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"),  "single")
        b.set(qn("w:sz"),   sz)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _set_cell_min_height(cell, cm: float) -> None:
    """Force a minimum row height so blank fill-in rows print cleanly."""
    tr = cell._tc.getparent()
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr"); tr.insert(0, trPr)
    height = OxmlElement("w:trHeight")
    # Word measures heights in twentieths of a point. 1cm ≈ 567 twips.
    height.set(qn("w:val"), str(int(cm * 567)))
    height.set(qn("w:hRule"), "atLeast")
    trPr.append(height)


def _set_cell_width_pct(cell, pct: int) -> None:
    """Set cell width as a percentage of the table width (1000 = 100%)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(pct * 50))   # ~50 = 1% in dxa heuristic; we use pct units
    tcW.set(qn("w:type"), "pct")
    # python-docx leaves stale width specs around; replace if present
    existing = tc_pr.find(qn("w:tcW"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(tcW)


def _add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_field(paragraph, instr: str) -> None:
    """Insert a Word field code (e.g. 'PAGE') in a paragraph as a run."""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " " + instr + " "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ──────────────────────────────────────────────────────────────────────
# Part 3 — Header-param extraction (parse the MD cover area)
# ──────────────────────────────────────────────────────────────────────

def extract_header_params(text: str) -> dict[str, str]:
    """Sniff the cover area of the MD for the values we need on cover +
    footer + forms. Robust to whitespace + formatting variations."""
    params: dict[str, str] = {
        "department":         "",
        "department_full":    "",
        "nit_number":         "",
        "issue_date":         "",
        "name_of_work":       "",
        "contact_officer":    "",
        "contact_email":      "",
        "ecv_rupees":         "",
        "ecv_cr":             "",
        "state_upper":        "GOVERNMENT OF ANDHRA PRADESH",
    }
    # Cover area = first ~60 lines (before the Standard Bidding Document heading)
    head = "\n".join(text.splitlines()[:80])

    # 1) Department acronym = the second H2 (first is "GOVERNMENT OF AP,")
    m = re.search(r"^##\s*\*\*([A-Z0-9]+)\*\*\s*$", head, re.M)
    if m:
        params["department"]      = m.group(1).strip()
        params["department_full"] = m.group(1).strip()

    # 2) NIT number — "NIT No : 100/PROC/APIIC/1/2026, Dt:08/05/2026"
    m = re.search(r"NIT No\s*:?\s*([^,*]+),\s*Dt\s*:?\s*([0-9]{1,2}[/\-][0-9]{1,2}[/\-][0-9]{2,4})", head)
    if m:
        params["nit_number"] = m.group(1).strip()
        params["issue_date"] = m.group(2).strip()

    # 3) Name of work — "**Name of the Work : Construction of …**"
    m = re.search(r"Name of the Work\s*:?\s*([^*]+?)(?:\*\*|$)", head, re.M)
    if m:
        params["name_of_work"] = m.group(1).strip().rstrip(".")

    # 4) Issued on
    m = re.search(r"ISSUED\s*ON\s*:?\s*([0-9]{1,2}[/\-][0-9]{1,2}[/\-][0-9]{2,4})", head, re.I)
    if m:
        params["issue_date"] = m.group(1).strip()

    # 5) ECV — pull from the NIT body table further down (line ~30-70)
    full_head = "\n".join(text.splitlines()[:150])
    m = re.search(r"Estimated Contract Value \(ECV\)\*?\*?\s*\|\s*\*?\*?Rs\.?([0-9,]+\.\d{2})", full_head)
    if m:
        params["ecv_rupees"] = m.group(1).replace(",", "")
    m = re.search(r"Rs\.?\s*([0-9]+(?:\.\d+)?)\s*Crore", full_head)
    if m:
        params["ecv_cr"] = m.group(1).strip()

    # 6) Contact officer + email — pull from NIT body
    m = re.search(r"Contact Person[^|]*\|\s*([^,(|]+),", full_head)
    if m:
        params["contact_officer"] = m.group(1).strip()
    m = re.search(r"\(([\w.+-]+@[\w-]+\.[\w.-]+)\)", full_head)
    if m:
        params["contact_email"] = m.group(1).strip()

    return params


def _drop_md_cover_lines(lines: list[str]) -> tuple[list[str], int]:
    """Find where the printed cover ends (first ``---`` after the
    `## ISSUED ON: …` heading) and return the lines AFTER that HR.
    Falls back to the first ``---`` if the header is non-standard.
    The DOCX cover page is rendered separately via _add_cover_page().
    """
    issued_idx = -1
    for i, ln in enumerate(lines[:30]):
        if ln.strip().lower().startswith("## **issued on"):
            issued_idx = i
            break
    start = issued_idx + 1 if issued_idx >= 0 else 0
    for i, ln in enumerate(lines[start:], start=start):
        if _HR_RE.match(ln.strip()):
            return lines[i+1:], i + 1
    return lines, 0


# ──────────────────────────────────────────────────────────────────────
# Part 4 — Cover page builder
# ──────────────────────────────────────────────────────────────────────

def _add_cover_page(doc: Document, params: dict[str, str]) -> None:
    """Build the cover page: seal placeholder, department name, two
    HRs, NIT number, name of work, issued-on date, page break.

    Layout matches the AGICL / ADCL bid documents the user supplied:
    centred mark, bold issuing-entity name in serif, two horizontal
    rules, NIT block, Name of the Work, Issued On, then a page break.
    """
    # ── 1. AP Government emblem placeholder ───────────────────────────
    # TODO: Replace with official AP seal PNG. Available from
    # ap.gov.in or the Press & Information Department. Until then we
    # render an 80×80 outlined box with "AP GOVT SEAL" text — visually
    # marks the intended placement so an officer can paste the real
    # PNG in Word using Insert → Picture.
    # Use a 3-column table with empty side cells so the centred 80×80
    # seal box doesn't stretch to the full page width.
    seal_table = doc.add_table(rows=1, cols=3)
    seal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    seal_table.autofit = False
    seal_table.columns[0].width = Cm(7.0)
    seal_table.columns[1].width = Cm(2.4)   # ~80 px at 96 DPI
    seal_table.columns[2].width = Cm(7.0)
    for ci in (0, 1, 2):
        c = seal_table.cell(0, ci)
        c.width = Cm(7.0 if ci != 1 else 2.4)
        # Hide borders on the side cells
        if ci != 1:
            tc_pr = c._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "nil")
                borders.append(b)
            tc_pr.append(borders)
    seal_cell = seal_table.cell(0, 1)
    _set_cell_borders(seal_cell, sz="8", color="808080")
    _set_cell_shading(seal_cell, "F2F2F2")
    _set_cell_min_height(seal_cell, 2.4)
    seal_p = seal_cell.paragraphs[0]
    seal_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    seal_p.paragraph_format.space_before = Pt(20)
    r = seal_p.add_run("AP GOVT\nSEAL")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ── 2. State header (small, centred) ──────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(params.get("state_upper", "GOVERNMENT OF ANDHRA PRADESH"))
    r.bold = True
    r.font.size = Pt(11)

    # ── 3. Department name (bold, centred, 14pt) ──────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept = params.get("department_full") or params.get("department") or "[DEPARTMENT]"
    r = p.add_run(dept)
    r.bold = True
    r.font.size = Pt(14)

    # ── 4. Two horizontal rules ───────────────────────────────────────
    _add_hr(doc)
    _add_hr(doc)

    # spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(60)

    # ── 5. NIT number (bold, underlined, 12pt, centred) ───────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    nit = params.get("nit_number") or "[NIT NUMBER]"
    issue = params.get("issue_date") or "[DATE]"
    r = p.add_run(f"NIT No. {nit}, Dt: {issue}")
    r.bold = True; r.underline = True; r.font.size = Pt(12)

    # ── 6. Name of the Work (bold, 11pt, centred, hanging indent) ─────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    name = params.get("name_of_work") or "[NAME OF WORK]"
    r = p.add_run("Name of the Work: ")
    r.bold = True; r.font.size = Pt(11)
    r = p.add_run(name)
    r.bold = True; r.font.size = Pt(11)

    # ── 7. Issued On date ─────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run(f"ISSUED ON: {issue}")
    r.bold = True; r.font.size = Pt(11)

    # ── 8. Page break ─────────────────────────────────────────────────
    _add_page_break(doc)


# ──────────────────────────────────────────────────────────────────────
# Part 5 — Page footer + typography
# ──────────────────────────────────────────────────────────────────────

def _add_page_footer(section, params: dict[str, str]) -> None:
    """Three-cell footer: 'Contractor' left, page-number centre,
    'Department, Officer' right. Times New Roman 9pt."""
    footer = section.footer
    # Clear any default empty paragraph
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    # Build a 1×3 borderless table inside the footer
    tbl = footer.add_table(rows=1, cols=3, width=Cm(17.0))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(5.5)
    tbl.columns[1].width = Cm(5.5)
    tbl.columns[2].width = Cm(6.0)

    # Left — "Contractor"
    left = tbl.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = left.add_run("Contractor")
    r.font.name = "Times New Roman"; r.font.size = Pt(9)

    # Centre — page number
    centre = tbl.cell(0, 1).paragraphs[0]
    centre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = centre.add_run("Page ")
    r.font.name = "Times New Roman"; r.font.size = Pt(9)
    _add_field(centre, "PAGE")
    r = centre.add_run(" of ")
    r.font.name = "Times New Roman"; r.font.size = Pt(9)
    _add_field(centre, "NUMPAGES")
    for run in centre.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    # Right — "Department, Officer"
    right = tbl.cell(0, 2).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dept    = params.get("department")      or params.get("department_full") or "[Department]"
    officer = params.get("contact_officer") or "[Officer]"
    r = right.add_run(f"{dept}, {officer}")
    r.font.name = "Times New Roman"; r.font.size = Pt(9)

    # Hide all cell borders inside the footer
    for cell in (tbl.cell(0, 0), tbl.cell(0, 1), tbl.cell(0, 2)):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            borders.append(b)
        tc_pr.append(borders)


def _apply_typography(doc: Document) -> None:
    """Set serif-body / sans-heading defaults that match real AP tenders.

    Body / Normal: Times New Roman 11pt, justified.
    Heading 1-6 : Arial 12-13pt bold (sans, distinguishable).
    """
    # Normal (body) — Times New Roman 11pt
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    # Force the East-Asian font slot too — Word otherwise falls back to Calibri
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    "Times New Roman")
    rFonts.set(qn("w:hAnsi"),    "Times New Roman")
    rFonts.set(qn("w:cs"),       "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if rFonts.getparent() is None:
        rPr.append(rFonts)

    # Heading 1-6 — Arial bold
    for lvl, sz in ((1, 14), (2, 13), (3, 12), (4, 11), (5, 11), (6, 11)):
        try:
            h = doc.styles[f"Heading {lvl}"]
            h.font.name = "Arial"
            h.font.size = Pt(sz)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            rPr = h.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), "Arial")
            rFonts.set(qn("w:hAnsi"), "Arial")
            rFonts.set(qn("w:cs"),    "Arial")
            if rFonts.getparent() is None:
                rPr.append(rFonts)
        except KeyError:
            pass


# ──────────────────────────────────────────────────────────────────────
# Part 6 — Forms classifier + render_form_section() + handlers
# ──────────────────────────────────────────────────────────────────────

# Map (lowercased title substring) → form-type identifier. Tested
# left-to-right; first match wins. Order matters — more specific
# substrings come before more generic ones.
_FORM_TYPE_DISPATCH: list[tuple[str, str]] = [
    ("statement-i — annual",            "statement_i_turnover"),
    ("statement-ii — similar works",    "statement_ii_similar_works"),
    ("statement-iii — critical equipment", "statement_iii_equipment"),
    ("statement-iv — key personnel",    "statement_iv_personnel"),
    ("statement-v — ongoing commitments","statement_v_ongoing"),
    ("statement-vi — liquid assets",    "statement_vi_liquid_assets"),
    ("performance bank guarantee",      "pbg_proforma"),
    ("advance payment guarantee",       "apg_proforma"),
    ("letter of acceptance",            "loa"),
    ("contract agreement",              "contract_agreement"),
    ("manufacturer's authorisation",    "maf"),
    # Generic: anything that is a "Declaration", "Certificate", or "Undertaking"
    ("declaration",                     "declaration"),
    ("certificate",                     "declaration"),
    ("undertaking",                     "declaration"),
    ("compliance certificate",          "declaration"),
    ("self-certification",              "declaration"),
]


def classify_form(title: str) -> str | None:
    t = (title or "").lower()
    for needle, ftype in _FORM_TYPE_DISPATCH:
        if needle in t:
            return ftype
    return None


def render_form_section(doc: Document, form_type: str, params: dict[str, Any],
                        *, title: str = "", body_md: str = "") -> bool:
    """Public API — render a single tender form structurally.

    Each form starts on a new page (page break inserted before the
    form heading). The call returns True if the form_type was
    recognised and rendered, False otherwise (so the caller can fall
    back to its default rendering).

    Args:
        doc:       the python-docx Document being built
        form_type: one of the identifiers from _FORM_TYPE_DISPATCH
        params:    dict with keys like nit_number, project_name,
                   department, contact_officer, ecv_rupees, ecv_cr,
                   issue_date, contact_email
        title:     the original markdown title (for fall-back display)
        body_md:   the original markdown body (used for declarations
                   so the actual declaration text is preserved)
    """
    # Page break + form title
    _add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper().strip())
    r.bold = True; r.underline = True; r.font.size = Pt(12)
    r.font.name = "Arial"

    # Dispatch
    if form_type == "statement_i_turnover":
        _render_statement_i(doc, params)
    elif form_type in ("statement_ii_similar_works", "statement_iii_equipment",
                        "statement_iv_personnel", "statement_v_ongoing",
                        "statement_vi_liquid_assets"):
        _render_generic_statement(doc, form_type, title, params)
    elif form_type == "pbg_proforma":
        _render_pbg_proforma(doc, params)
    elif form_type == "apg_proforma":
        _render_apg_proforma(doc, params)
    elif form_type == "loa":
        _render_loa(doc, params)
    elif form_type == "contract_agreement":
        _render_contract_agreement(doc, params)
    elif form_type == "maf":
        _render_maf(doc, params)
    elif form_type == "declaration":
        _render_declaration(doc, title, body_md, params)
    else:
        return False
    return True


# ── Sub-helpers used by form renderers ─────────────────────────────────

def _form_info_block(doc: Document, params: dict[str, Any]) -> None:
    """Two-column 'Bidder Name / Tender No' info block at the top of forms."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for col in tbl.columns:
        col.width = Cm(8.5)
    nit = params.get("nit_number") or ""
    issue = params.get("issue_date") or ""
    pairs = [
        ("Bidder Name", "_" * 30),
        ("Tender No",   f"{nit}  Dt. {issue}"),
        ("Name of Work", params.get("name_of_work") or "_" * 30),
        ("Date of Submission", "_" * 20),
    ]
    # Only need 2 rows of 2 pairs each — flatten to 4 cells
    flat = pairs
    for i, (k, v) in enumerate(flat):
        cell = tbl.cell(i // 2, i % 2)
        _set_cell_borders(cell, sides=("bottom",), sz="6", color="808080")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(f"{k}: ")
        r.bold = True
        p.add_run(str(v))
    doc.add_paragraph()  # spacer


def _signature_block(doc: Document, *, with_ca: bool = False) -> None:
    """Two-column signature block. If with_ca=True, right column is
    'CA Counter-signature' (used for Statements I-VI); otherwise right
    column is 'Witness'."""
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in tbl.columns: col.width = Cm(8.5)

    left_lines = [
        ("Signature",    "_" * 30),
        ("Name",         "_" * 30),
        ("Designation",  "_" * 30),
        ("Date",         "_" * 30),
        ("Place",        "_" * 30),
        ("",             "(Firm Seal)"),
    ]
    if with_ca:
        right_lines = [
            ("CA Counter-signature", "_" * 25),
            ("Name",                 "_" * 30),
            ("Membership No",        "_" * 25),
            ("Firm Reg No",          "_" * 25),
            ("Date",                 "_" * 30),
            ("",                     "(CA Firm Seal)"),
        ]
    else:
        right_lines = [
            ("Witness Signature", "_" * 25),
            ("Name",              "_" * 30),
            ("Address",           "_" * 30),
            ("",                  "_" * 30),
            ("Date",              "_" * 30),
            ("",                  ""),
        ]

    def _fill_cell(cell, lines):
        cell.paragraphs[0].text = ""    # clear default
        for k, v in lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            if k:
                r = p.add_run(f"{k}: "); r.bold = True
            p.add_run(str(v))
        # remove the empty paragraph[0] we cleared
        if not cell.paragraphs[0].text:
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    _fill_cell(tbl.cell(0, 0), left_lines)
    _fill_cell(tbl.cell(0, 1), right_lines)


def _section_header_box(doc: Document, line1: str, line2: str = "") -> None:
    """Full-width header box used at the top of Statement forms."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_borders(cell, sz="8", color="000000")
    _set_cell_shading(cell, "EFEFEF")
    _set_cell_min_height(cell, 1.6)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(line1.upper())
    r.bold = True; r.font.size = Pt(13); r.font.name = "Arial"
    if line2:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        r = p2.add_run(line2.upper())
        r.bold = True; r.font.size = Pt(11); r.font.name = "Arial"
    doc.add_paragraph()  # spacer


# ── Form: Statement-I — Annual Turnover (5 years) ───────────────────────

def _render_statement_i(doc: Document, params: dict[str, Any]) -> None:
    _section_header_box(doc, "Statement — I",
                        "Annual Financial Turnover (last 5 financial years)")
    _form_info_block(doc, params)

    # Derive the 5 financial years ending the most recent FY
    issue = params.get("issue_date") or ""
    try:
        d = datetime.strptime(issue, "%d/%m/%Y")
        # Most recent completed FY is (d.year - 1) to d.year if before April,
        # else d.year to d.year + 1. Here we want LAST 5 — work backwards.
        last_fy_end = d.year if d.month >= 4 else d.year - 1
    except Exception:
        last_fy_end = datetime.utcnow().year - 1
    fys = [f"{last_fy_end - i - 1}-{str(last_fy_end - i)[-2:]}"
           for i in range(5)]

    # Data table — 6 cols, 6 rows (header + 5 data)
    tbl = doc.add_table(rows=6, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["S.No", "Financial Year", "Turnover (Rs. Cr)",
               "Audited / Provisional", "Auditor Reference"]
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        _set_cell_borders(c, sz="8", color="000000")
        _set_cell_shading(c, "D9E1F2")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10)
        r.font.name = "Times New Roman"
    for ri in range(1, 6):
        for ci in range(5):
            c = tbl.cell(ri, ci)
            _set_cell_borders(c, sz="8", color="000000")
            _set_cell_min_height(c, 1.0)
            p = c.paragraphs[0]
            r = p.add_run(str(ri) if ci == 0 else
                          (fys[ri - 1] if ci == 1 else ""))
            r.font.name = "Times New Roman"; r.font.size = Pt(10)

    # Summary line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Average Annual Turnover (5 years): Rs. ")
    r.bold = True
    p.add_run("_" * 22)
    r = p.add_run("  Crore")
    r.bold = True

    # Certification block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Certification: ")
    r.bold = True
    p.add_run(
        "I/We certify that the above figures are extracted from the "
        "audited Profit & Loss Statements of the Bidder for the "
        "respective financial years. Copies of the audited statements "
        "(or, where audit is pending, provisional statements with the "
        "Statutory Auditor's certificate) are enclosed as supporting "
        "documents."
    )

    # Two-column signature block
    _signature_block(doc, with_ca=True)


# ── Form: Statement-II / III / IV / V / VI — generic 5-row form ─────────

_GENERIC_STATEMENT_HEADERS: dict[str, list[str]] = {
    "statement_ii_similar_works": [
        "S.No", "Project Name + Client", "Contract Value (Rs. Cr)",
        "Year of Completion", "Scope (similar to this Works)",
        "Completion Certificate Reference",
    ],
    "statement_iii_equipment": [
        "S.No", "Equipment Description", "Make / Model", "Capacity",
        "Owned / Leased / To be procured", "Quantity",
    ],
    "statement_iv_personnel": [
        "S.No", "Position", "Name (proposed)",
        "Years of Relevant Experience", "Qualification",
        "Years with Bidder",
    ],
    "statement_v_ongoing": [
        "S.No", "Project Name + Client", "Contract Value (Rs. Cr)",
        "Value of Outstanding Work (Rs. Cr)",
        "Stipulated Date of Completion", "Anticipated Slippage",
    ],
    "statement_vi_liquid_assets": [
        "S.No", "Source", "Type (Bank Balance / FD / OD limit / Solvency)",
        "Amount (Rs. Cr)", "Bank Name", "Reference",
    ],
}


def _render_generic_statement(doc: Document, form_type: str,
                              title: str, params: dict[str, Any]) -> None:
    headers = _GENERIC_STATEMENT_HEADERS.get(form_type, [
        "S.No", "Field 1", "Field 2", "Field 3", "Field 4", "Field 5"])
    # Use the title sans-prefix as the box header
    box = title
    if "—" in title:
        head, tail = title.split("—", 1)
        box = head.strip()
        sub = tail.strip()
    else:
        sub = ""
    _section_header_box(doc, box, sub)
    _form_info_block(doc, params)

    n_cols = len(headers)
    tbl = doc.add_table(rows=6, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        _set_cell_borders(c, sz="8", color="000000")
        _set_cell_shading(c, "D9E1F2")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10)
        r.font.name = "Times New Roman"
    for ri in range(1, 6):
        for ci in range(n_cols):
            c = tbl.cell(ri, ci)
            _set_cell_borders(c, sz="8", color="000000")
            _set_cell_min_height(c, 1.0)
            if ci == 0:
                p = c.paragraphs[0]
                r = p.add_run(str(ri))
                r.font.name = "Times New Roman"; r.font.size = Pt(10)

    _signature_block(doc, with_ca=True)


# ── Form: Performance Bank Guarantee Proforma ──────────────────────────

def _render_pbg_proforma(doc: Document, params: dict[str, Any]) -> None:
    nit = params.get("nit_number") or "[NIT NUMBER]"
    name = params.get("name_of_work") or params.get("project_name") or "[NAME OF WORK]"
    dept = params.get("department") or "[EMPLOYER]"
    ecv_cr = params.get("ecv_cr") or ""

    # Bank letterhead box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_borders(cell, sz="6", color="808080")
    _set_cell_shading(cell, "F8F8F8")
    _set_cell_min_height(cell, 2.4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("[ BANK LETTERHEAD ]"); r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80); r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("(Issued on the letterhead of a Nationalised / Scheduled Commercial Bank "
                    "on non-judicial stamp paper of value as applicable)")
    r2.italic = True; r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r2.font.size = Pt(9)

    doc.add_paragraph()

    # Reference / date line
    ref = doc.add_paragraph()
    ref.add_run("Bank Guarantee No.: ").bold = True
    ref.add_run("_" * 30)
    ref.add_run("    ")
    ref.add_run("Date: ").bold = True
    ref.add_run("_" * 22)

    # To-block
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    p.add_run("To,")
    p = doc.add_paragraph()
    p.add_run(f"The {dept},").bold = True
    p = doc.add_paragraph(); p.add_run("_" * 60)
    p = doc.add_paragraph(); p.add_run("_" * 60)

    # Body
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_before = Pt(14)
    body1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body1.add_run("WHEREAS ").bold = True
    body1.add_run("M/s ")
    body1.add_run("_" * 30)
    body1.add_run(" (hereinafter called 'the Contractor') has been awarded "
                  f"the contract for the work of ")
    r = body1.add_run(name); r.bold = True
    body1.add_run(f" under Letter of Acceptance (LoA) No. ")
    body1.add_run("_" * 16)
    body1.add_run(" dated ")
    body1.add_run("_" * 16)
    body1.add_run(f", read with NIT No. {nit};")

    body2 = doc.add_paragraph()
    body2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body2.add_run("AND WHEREAS ").bold = True
    body2.add_run("the Contractor is required to furnish to the Employer a Performance "
                  "Security in the form of a Bank Guarantee equal to ")
    r = body2.add_run("10% of the contract value"); r.bold = True
    body2.add_run(" for the due performance of the contract;")

    # Prominent amount line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GUARANTEE AMOUNT: Rs. ")
    r.bold = True; r.font.size = Pt(12)
    p.add_run("_" * 24)
    r = p.add_run("/-")
    r.bold = True; r.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("(Rupees ")
    r.bold = True; r.font.size = Pt(11)
    p2.add_run("_" * 50)
    p2.add_run(" only)")

    # Validity + signature
    body3 = doc.add_paragraph()
    body3.paragraph_format.space_before = Pt(12)
    body3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body3.add_run(
        "The Bank further agrees that this Guarantee shall be valid until "
        "60 (sixty) days after the expiry of the Defects Liability Period, "
        "viz. until "
    )
    body3.add_run("_" * 16)
    body3.add_run(
        ". Notwithstanding anything to the contrary, this Bank Guarantee "
        "shall stand discharged only upon the Employer issuing a written "
        "certificate that the Contractor has duly performed all his "
        "obligations under the contract."
    )

    # Signature block (bank-side)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.add_run("For and on behalf of [BANK NAME]").bold = True

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for col in tbl.columns: col.width = Cm(8.5)

    lines_left = [
        ("Authorised Signatory", "_" * 30),
        ("Name",                 "_" * 30),
        ("Designation",          "_" * 30),
        ("Date",                 "_" * 30),
    ]
    lines_right = [
        ("Bank Stamp / Seal", ""),
    ]

    cell_l = tbl.cell(0, 0); cell_l.paragraphs[0].text = ""
    for k, v in lines_left:
        p = cell_l.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{k}: "); r.bold = True
        p.add_run(str(v))
    if not cell_l.paragraphs[0].text:
        cell_l.paragraphs[0]._element.getparent().remove(cell_l.paragraphs[0]._element)

    cell_r = tbl.cell(0, 1); cell_r.paragraphs[0].text = ""
    p = cell_r.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Bank Stamp / Seal:"); r.bold = True
    seal_box = doc.add_table(rows=1, cols=1)
    seal_box.alignment = WD_TABLE_ALIGNMENT.LEFT
    seal_cell = seal_box.cell(0, 0)
    _set_cell_borders(seal_cell, sz="8", color="000000")
    _set_cell_min_height(seal_cell, 2.5)
    cell_r.add_paragraph()  # spacer; the seal table sits below the signature block

    # Note: the boxed seal is a sibling of the signature 2-col table in
    # the doc body, not inside the right cell — Word renders nested
    # tables inconsistently across viewers, so we put the seal table
    # immediately after.

    _ = ecv_cr  # quietly unused; kept available if a future amount-line uses it


# ── Form: Advance Payment Guarantee Proforma (similar to PBG, smaller) ──

def _render_apg_proforma(doc: Document, params: dict[str, Any]) -> None:
    name = params.get("name_of_work") or "[NAME OF WORK]"
    dept = params.get("department") or "[EMPLOYER]"

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_borders(cell, sz="6", color="808080")
    _set_cell_shading(cell, "F8F8F8")
    _set_cell_min_height(cell, 2.4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("[ BANK LETTERHEAD ]"); r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80); r.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph().add_run("To,")
    p = doc.add_paragraph()
    p.add_run(f"The {dept}").bold = True
    p.add_run("  ____________________________________________")

    body = doc.add_paragraph()
    body.paragraph_format.space_before = Pt(14)
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run("WHEREAS ").bold = True
    body.add_run("M/s "); body.add_run("_" * 28)
    body.add_run(" (hereinafter called 'the Contractor') has been awarded the contract for the "
                 "work of ")
    r = body.add_run(name); r.bold = True
    body.add_run(", and is entitled to receive a mobilisation advance of ")
    r = body.add_run("10% of the contract value"); r.bold = True
    body.add_run(" against this Guarantee;")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GUARANTEE AMOUNT: Rs. ")
    r.bold = True; r.font.size = Pt(12)
    p.add_run("_" * 24)
    r = p.add_run("/-")
    r.bold = True; r.font.size = Pt(12)

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_before = Pt(12)
    body3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body3.add_run(
        "This Guarantee shall be REDUCED proportionately on each running-"
        "bill recovery of the advance, and shall stand DISCHARGED upon "
        "full recovery."
    )

    # Signature
    doc.add_paragraph().paragraph_format.space_before = Pt(18)
    p = doc.add_paragraph()
    p.add_run("For and on behalf of [BANK NAME]").bold = True
    for k in ("Authorised Signatory", "Name", "Designation", "Date"):
        p = doc.add_paragraph()
        r = p.add_run(f"{k}: "); r.bold = True
        p.add_run("_" * 36)


# ── Form: Letter of Acceptance ─────────────────────────────────────────

def _render_loa(doc: Document, params: dict[str, Any]) -> None:
    nit  = params.get("nit_number") or "[NIT NUMBER]"
    name = params.get("name_of_work") or "[NAME OF WORK]"
    dept = params.get("department_full") or params.get("department") or "[Department]"

    # Department letterhead area
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(dept.upper())
    r.bold = True; r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(params.get("state_upper", "GOVERNMENT OF ANDHRA PRADESH"))
    r.bold = True; r.font.size = Pt(10)

    _add_hr(doc)
    doc.add_paragraph()

    # Letter no + date (right aligned in real docs, here left for simplicity)
    p = doc.add_paragraph()
    r = p.add_run("Letter No.: "); r.bold = True
    p.add_run("_" * 24)
    p.add_run("    ")
    r = p.add_run("Date: "); r.bold = True
    p.add_run("_" * 18)

    # To
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
    p.add_run("To,")
    p = doc.add_paragraph(); p.add_run("M/s ").bold = True; p.add_run("_" * 50)
    p = doc.add_paragraph(); p.add_run("_" * 60)
    p = doc.add_paragraph(); p.add_run("_" * 60)

    # Subject
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Sub: "); r.bold = True
    p.add_run(f"Award of Contract for {name} — Letter of Acceptance")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Ref: "); r.bold = True
    p.add_run(
        f"(i) NIT No. {nit};  (ii) Your Bid dated _____________; "
        "(iii) Technical evaluation dated _____________; "
        "(iv) Financial evaluation dated _____________."
    )

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
    p.add_run("Sir,")

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Cm(0.75)
    body.add_run("1.  With reference to the above, the ")
    body.add_run(dept).bold = True
    body.add_run(
        " is pleased to ACCEPT your bid for the above-named work for a "
        "contract value of "
    )
    r = body.add_run("Rs. " + ("_" * 14) + " Crore (Rupees " + ("_" * 36) + " only)")
    r.bold = True
    body.add_run(", at a percentage of "); body.add_run("____________")
    body.add_run(" with respect to the Estimated Contract Value (ECV) of Rs. ")
    body.add_run(params.get("ecv_cr") or "____")
    body.add_run(" Crore, on the terms and conditions of the bidding document.")

    body2 = doc.add_paragraph()
    body2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body2.paragraph_format.first_line_indent = Cm(0.75)
    body2.add_run("2.  You are required to: ")
    body2.add_run("(a) Furnish a Performance Bank Guarantee for ")
    r = body2.add_run("10% of the contract value"); r.bold = True
    body2.add_run(", valid until 60 days after the expiry of the Defects "
                  "Liability Period, within ")
    body2.add_run("________"); body2.add_run(" days from the date of this LoA; "
                  "(b) Sign and return the Contract Agreement within ")
    r = body2.add_run("14 (fourteen) days"); r.bold = True
    body2.add_run("; (c) Mobilise to Site within "); body2.add_run("________")
    body2.add_run(" days of signing the Contract Agreement.")

    body3 = doc.add_paragraph()
    body3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body3.paragraph_format.first_line_indent = Cm(0.75)
    body3.add_run("3.  ")
    r = body3.add_run("Failure to submit the PBG OR sign the Contract Agreement "
                      "within the stipulated time")
    r.bold = True
    body3.add_run(" shall constitute sufficient grounds for the annulment of "
                  "this Award and forfeiture of the Bid Security per ITB §42.2.")

    # Sign-off
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Yours faithfully,")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("_" * 30)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("(Authorised Signatory)").bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("[Office Seal]").italic = True


# ── Form: Contract Agreement (compact stub) ────────────────────────────

def _render_contract_agreement(doc: Document, params: dict[str, Any]) -> None:
    name = params.get("name_of_work") or "[NAME OF WORK]"
    dept = params.get("department_full") or params.get("department") or "[Department]"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("THIS CONTRACT AGREEMENT ").bold = True
    p.add_run("is made on this ")
    p.add_run("_____ "); p.add_run("day of "); p.add_run("____________ ")
    p.add_run("Two Thousand "); p.add_run("____________ ")
    p.add_run("BETWEEN ")
    r = p.add_run(dept); r.bold = True
    p.add_run(" (hereinafter called 'the Employer') of the ONE PART, AND ")
    r = p.add_run("M/s ____________________"); r.bold = True
    p.add_run(" (hereinafter called 'the Contractor') of the OTHER PART.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("WHEREAS ").bold = True
    p.add_run(f"the Employer desires that the work known as ")
    r = p.add_run(name); r.bold = True
    p.add_run(" should be executed by the Contractor and has accepted the "
              "Contractor's bid for the same.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("NOW THIS AGREEMENT WITNESSETH "); r.bold = True
    p.add_run("as follows: ")
    p.add_run("(1) The Contract Documents shall be deemed to form and be read "
              "and construed as part of this Agreement; (2) The Contractor "
              "shall execute and complete the Works in conformity in all "
              "respects with the provisions of the Contract; (3) The Employer "
              "shall pay the Contractor the Contract Price in consideration of "
              "the execution and completion of the Works.")

    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    tbl = doc.add_table(rows=1, cols=2)
    for col in tbl.columns: col.width = Cm(8.5)

    cell_l = tbl.cell(0, 0); cell_l.paragraphs[0].text = ""
    for k in ("For the Employer", "Signature", "Name", "Designation", "Date"):
        p = cell_l.add_paragraph()
        if k == "For the Employer":
            r = p.add_run(k); r.bold = True
        else:
            r = p.add_run(f"{k}: "); r.bold = True; p.add_run("_" * 30)
    if not cell_l.paragraphs[0].text:
        cell_l.paragraphs[0]._element.getparent().remove(cell_l.paragraphs[0]._element)

    cell_r = tbl.cell(0, 1); cell_r.paragraphs[0].text = ""
    for k in ("For the Contractor", "Signature", "Name", "Designation", "Date"):
        p = cell_r.add_paragraph()
        if k == "For the Contractor":
            r = p.add_run(k); r.bold = True
        else:
            r = p.add_run(f"{k}: "); r.bold = True; p.add_run("_" * 30)
    if not cell_r.paragraphs[0].text:
        cell_r.paragraphs[0]._element.getparent().remove(cell_r.paragraphs[0]._element)


# ── Form: Manufacturer's Authorisation Form ────────────────────────────

def _render_maf(doc: Document, params: dict[str, Any]) -> None:
    nit  = params.get("nit_number") or "[NIT NUMBER]"
    name = params.get("name_of_work") or "[NAME OF WORK]"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("To,").bold = True
    p = doc.add_paragraph(); p.add_run("_" * 60)
    p = doc.add_paragraph(); p.add_run("_" * 60)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Sub: "); r.bold = True
    p.add_run(f"Manufacturer's Authorisation under NIT No. {nit} for the work of ")
    r = p.add_run(name); r.bold = True
    p.add_run(".")

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(
        "We, "); body.add_run("M/s ____________________________________ "
                              "(Manufacturer's Name and Address)")
    body.add_run(", hereby authorise ")
    body.add_run("M/s ____________________________________ ")
    body.add_run("(hereinafter the 'Bidder') to submit a bid in connection "
                 "with the above-referenced tender, the supply being of the "
                 "products manufactured by us. We hereby extend our full "
                 "guarantee and warranty in accordance with the General "
                 "Conditions of Contract, with respect to the products offered "
                 "by the Bidder.")

    doc.add_paragraph().paragraph_format.space_before = Pt(18)
    for k in ("Signature of authorised signatory of Manufacturer",
              "Name", "Designation", "Date"):
        p = doc.add_paragraph()
        r = p.add_run(f"{k}: "); r.bold = True
        p.add_run("_" * 36)


# ── Form: Generic Declaration / Certificate / Undertaking ───────────────

def _render_declaration(doc: Document, title: str,
                        body_md: str, params: dict[str, Any]) -> None:
    """For declarations / certificates / undertakings — render a
    centred bold heading, the declaration text from the markdown body
    (stripped of the leak placeholders), self-addressed blanks where
    needed, and a signature block at the bottom."""
    # Body — break the markdown into paragraphs at sentence boundaries.
    # Replace common self-addressed slots with underline blanks.
    text = body_md or ""
    # The markdown body of declarations often arrives with trailing
    # blanks like "Bidder Name: _________________________." We keep
    # those untouched. We just split on multiple spaces to clean up.
    text = re.sub(r"\s+", " ", text).strip()
    # Soft-wrap into ~2 paragraphs per declaration (split at each ". ")
    sentences = re.split(r"(?<=[.!?])\s+", text)
    if not sentences:
        sentences = [text]

    # If the declaration looks like a list of numbered points, preserve
    # the numbers — they were re-collapsed by the whitespace normalisation
    # above only if no `1.` / `2.` markers were present.
    if any(re.match(r"^\d+\.\s", s) for s in sentences):
        # Keep one paragraph per numbered point.
        for s in sentences:
            if not s: continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_runs(p, s)
    else:
        # Drop "I/We …" style preamble into a self-addressed paragraph
        # with a blank for the bidder name.
        if sentences and sentences[0].lower().startswith(("i ", "i,", "i/we", "we ", "we,")):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run("I / We, "); r.bold = True
            p.add_run("_" * 38)
            p.add_run(", hereinafter called the Bidder, declare as follows:")
            sentences = sentences[1:]
        # Body
        if sentences:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_runs(p, " ".join(sentences))

    # Signature block (no CA counter-signature for declarations)
    _signature_block(doc, with_ca=False)


# ──────────────────────────────────────────────────────────────────────
# Part 7 — NIT body table styler + general table renderer
# ──────────────────────────────────────────────────────────────────────

# Labels in the NIT body whose right-column value should be bold
# (the figures / dates a procurement officer most often checks first).
_NIT_BOLD_LABELS = {
    "estimated contract value (ecv)",
    "period of completion of work",
    "period of defect liability period (dlp)",
    "bid validity",
    "bid security (emd)",
    "performance security (pbg)",
    "transaction fee",
    "bid document downloading start date",
    "bid document downloading close date",
    "pre-bid meeting date",
    "bid submission due date and time",
    "opening of technical bid",
    "opening of financial bid",
    "tender number",
    "tender subject",
    "department",
    "form of contract",
    "tender type",
    "eligible class of bidders and additional references",
    "category of registration",
}


def _add_table(doc: Document, rows: list[list[str]],
               *, is_nit_body: bool = False) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    # NIT body tables: 40/60 column widths
    if is_nit_body and n_cols == 2:
        tbl.columns[0].width = Cm(6.8)
        tbl.columns[1].width = Cm(10.2)

    for ri, row in enumerate(rows):
        cells = tbl.rows[ri].cells
        for ci in range(n_cols):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if is_nit_body and n_cols == 2:
                cell.width = Cm(6.8 if ci == 0 else 10.2)

            chunks = txt.split("<br>")
            cell.text = ""
            for k, chunk in enumerate(chunks):
                p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT
                _add_runs(p, chunk.strip())
                # NIT body: bold the value column when label matches
                if is_nit_body and ci == 1 and ri > 0:
                    label = re.sub(r"\*+", "", row[0]).strip().lower()
                    if label in _NIT_BOLD_LABELS:
                        for run in p.runs:
                            run.bold = True
            if ri == 0:
                _set_cell_shading(cell, "D9E1F2")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            _set_cell_borders(cell)


# ──────────────────────────────────────────────────────────────────────
# Part 8 — md_to_docx() — the top-level entry
# ──────────────────────────────────────────────────────────────────────

def md_to_docx(md_path: str | Path, docx_path: str | Path) -> str:
    md_path  = Path(md_path)
    docx_path = Path(docx_path)
    text = md_path.read_text(encoding="utf-8")

    doc = Document()

    # Page setup — A4 portrait
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.0)

    # Typography
    _apply_typography(doc)

    # Header params from the cover area
    params = extract_header_params(text)

    # Cover page (replaces the printed cover lines from the MD)
    _add_cover_page(doc, params)

    # Page footer — runs on every page including the cover
    _add_page_footer(section, params)

    # Drop the printed cover lines from the MD body before parsing
    raw_lines = text.splitlines()
    body_lines, _drop_count = _drop_md_cover_lines(raw_lines)

    # State machine: are we inside Section IV ("Bidding Forms")?
    in_section_iv = False
    section_iv_re = re.compile(r"^##\s*\*\*?\s*Section\s+IV\b", re.I)
    new_section_re = re.compile(r"^##\s*\*\*?\s*(?:Section\s+[IVX]+|PART\s+\d|NOTICE INVITING TENDER)", re.I)

    # NIT body detection: any 2-col table that follows the
    # "Notice Inviting Tender (NIT)" heading until the next H2
    in_nit_body = False
    nit_re = re.compile(r"^##\s*\*\*?\s*Notice Inviting Tender", re.I)

    i, n = 0, len(body_lines)
    while i < n:
        line = body_lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # HTML-comment line — skip; these are machine-readable markers
        # for the substitution layer (BDS_REF / PCC_REF / skeleton
        # provenance comments) and must not appear in the rendered DOCX.
        if _HTML_COMMENT_LINE_RE.match(stripped):
            i += 1
            continue

        # Section state transitions
        if _HEADING_RE.match(stripped):
            if section_iv_re.search(stripped):
                in_section_iv = True; in_nit_body = False
            elif nit_re.search(stripped):
                in_nit_body = True; in_section_iv = False
            elif new_section_re.search(stripped):
                in_section_iv = False; in_nit_body = False

        # Horizontal rule
        if _HR_RE.match(stripped):
            _add_hr(doc); i += 1; continue

        # ATX heading
        m = _HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            text_h = m.group(2).strip()
            # Strip the redundant "**…**" wrapping the drafter emits
            text_h = re.sub(r"^\*+|\*+$", "", text_h).strip()
            h = doc.add_heading(level=level if level <= 6 else 6)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_runs(h, text_h)
            i += 1; continue

        # Pipe table
        if _is_table_row(line) and i + 1 < n and _TABLE_SEP_RE.match(body_lines[i+1]):
            header = _strip_pipes(line)
            body_rows: list[list[str]] = []
            j = i + 2
            while j < n and _is_table_row(body_lines[j]):
                body_rows.append(_strip_pipes(body_lines[j]))
                j += 1

            # ── Forms section: dispatch each row to render_form_section ──
            if in_section_iv and len(header) >= 2:
                # The drafter emits Section IV as a 2-col "Form | Form Spec"
                # table. Each body row is one form. We ignore the header
                # row's label and dispatch on the title text.
                for row in body_rows:
                    if len(row) < 1: continue
                    title    = row[0].strip()
                    body_blob = row[1].strip() if len(row) > 1 else ""
                    ftype = classify_form(title)
                    if ftype is None:
                        # Unknown form — render as a one-row table so we
                        # don't lose content
                        _add_table(doc, [[title, body_blob]])
                    else:
                        render_form_section(
                            doc, ftype, params,
                            title=title, body_md=body_blob,
                        )
                i = j
                continue

            # ── NIT body: 2-col with bold key values + 40/60 widths ──
            _add_table(doc, [header] + body_rows,
                       is_nit_body=(in_nit_body and len(header) == 2))
            i = j
            continue

        # Bullet list
        if _BULLET_RE.match(stripped):
            while i < n and (m2 := _BULLET_RE.match(body_lines[i].strip())):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, m2.group(1))
                i += 1
            continue

        # Numbered list
        if _NUMLIST_RE.match(stripped):
            while i < n and (m2 := _NUMLIST_RE.match(body_lines[i].strip())):
                p = doc.add_paragraph(style="List Number")
                _add_runs(p, m2.group(1))
                i += 1
            continue

        # Blockquote — italic, indented
        if stripped.startswith(">"):
            buf: list[str] = []
            while i < n and body_lines[i].strip().startswith(">"):
                buf.append(body_lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            _add_runs(p, " ".join(buf))
            for r in p.runs:
                r.italic = True
            continue

        # Plain paragraph
        buf = [line]
        i += 1
        while i < n and body_lines[i].strip() and not (
            _HEADING_RE.match(body_lines[i].strip())
            or _is_table_row(body_lines[i])
            or _BULLET_RE.match(body_lines[i].strip())
            or _NUMLIST_RE.match(body_lines[i].strip())
            or _HR_RE.match(body_lines[i].strip())
            or body_lines[i].strip().startswith(">")
        ):
            buf.append(body_lines[i])
            i += 1
        para_text = " ".join(s.strip() for s in buf).strip()
        # Strip any inline HTML-comment markers that slipped into a
        # paragraph (e.g. when an inline `<!-- {{BDS_REF: …}} -->`
        # markers appears mid-line rather than on its own line).
        para_text = _HTML_COMMENT_INLINE_RE.sub("", para_text).strip()
        if para_text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_runs(p, para_text)

    doc.save(str(docx_path))
    return str(docx_path.resolve())


# ──────────────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("usage: python -m modules.draft_export.md_to_docx <md> [<docx>]")
        sys.exit(2)
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2]) if len(sys.argv) >= 3 else src.with_suffix(".docx")
    out = md_to_docx(src, dst)
    print(f"wrote {out}")
