"""
M1.7 — Artifact rendering (DOCX + PDF + BoQ XLSX + Eligibility DOCX).

Called at AUTHORITY publish gate to produce the final BID DOCUMENT
package. All artifacts are saved to /tmp/m1_artifacts/{draft_id}/v{version}/.

Reuses python-docx + reportlab (already installed for ComparativeStatement).
BoQ XLSX uses openpyxl when available, falls back to CSV if not.
"""
from __future__ import annotations

import csv
import json
import os
from pathlib import Path

from .schemas import TenderDraftState


def _artifact_dir(state: TenderDraftState) -> Path:
    p = Path(f"/tmp/m1_artifacts/{state.draft_id}/v{state.version}")
    p.mkdir(parents=True, exist_ok=True)
    return p


# ─── DOCX: Bid Document (full 15-section eGP format) ─────────────────


def render_bid_document_docx(state: TenderDraftState, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Header
    t = doc.add_heading("BID DOCUMENT", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Government of Andhra Pradesh · eGP Portal")
    r.italic = True
    r.font.size = Pt(11)

    # ─── Section 1: Current Tender Details ──────────────────────────
    doc.add_heading("1. Current Tender Details", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    _row(tbl, "Tender ID", state.tender_id or "(to be assigned)")
    _row(tbl, "Tender Notice Number", state.tender_notice_number or "")
    _row(tbl, "Name of Work", state.enquiry_particulars.name_of_work)
    _row(tbl, "Tender Category", state.classification.tender_category.value)
    _row(tbl, "Tender Type", state.classification.tender_type.value)
    _row(tbl, "Estimated Contract Value", f"₹ {state.financial.estimated_contract_value_inr:,}")
    _row(tbl, "Submission Closing Date", state.dates.closing_date)
    _row(tbl, "Tender Evaluation Type", state.evaluation.evaluation_type.value)

    # ─── Section 2: Enquiry Particulars ─────────────────────────────
    doc.add_heading("2. Enquiry Particulars", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    _row(tbl, "Department Name", state.enquiry_particulars.department_name)
    _row(tbl, "Circle / Division", state.enquiry_particulars.circle_division)
    _row(tbl, "IFB No / Tender Notice Number", state.tender_notice_number or "")
    _row(tbl, "Name of Project", state.enquiry_particulars.name_of_project)
    _row(tbl, "Name of Work", state.enquiry_particulars.name_of_work)
    _row(tbl, "Estimated Contract Value (INR)",
         f"₹ {state.financial.estimated_contract_value_inr:,}\n"
         f"({state.financial.estimated_contract_value_words})")
    _row(tbl, "Period of Completion (months)", str(state.financial.period_of_completion_months))
    _row(tbl, "Type of Work", state.classification.type_of_work)
    _row(tbl, "Bidding Type", state.classification.bidding_type.value)
    _row(tbl, "Bid Call (Numbers)", str(state.classification.bid_call_numbers))
    _row(tbl, "Tender Category", state.classification.tender_category.value)
    _row(tbl, "Currency Type", f"({state.financial.currency_type.value})")
    _row(tbl, "Default Currency", state.financial.default_currency)
    _row(tbl, "Evaluation Type", state.evaluation.evaluation_type.value)
    _row(tbl, "Evaluation Criteria", state.evaluation.evaluation_criteria.value)
    _row(tbl, "Form of Contract", state.classification.form_of_contract.value)
    _row(tbl, "Consortium / Joint Venture", state.classification.consortium_joint_venture.value)

    # ─── Section 3: Transaction Fee Details ─────────────────────────
    doc.add_heading("3. Transaction Fee Details", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    _row(tbl,
         f"Transaction Fee Payable to '{state.financial.transaction_fee_payable_to}'\n"
         f"(As per {state.financial.transaction_fee_go_reference})",
         f"{state.financial.transaction_fee_inr} (INR)")

    # ─── Section 4: Tender Dates ────────────────────────────────────
    doc.add_heading("4. Tender Dates", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    _row(tbl, "Start Date & Time", state.dates.start_date)
    _row(tbl, "End Date & Time", state.dates.end_date)
    _row(tbl, "Closing Date & Time", state.dates.closing_date)
    _row(tbl, "Bid Validity Period (Days)", str(state.financial.bid_validity_days))
    _row(tbl, "Display Rank", state.evaluation.display_rank.value)

    # ─── Section 5: Tender Inviting Authority ───────────────────────
    doc.add_heading("5. Tender Inviting Authority Particulars", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    _row(tbl, "Officer Inviting Bids", state.enquiry_particulars.officer_inviting_bids)
    _row(tbl, "Bid Opening Authority", state.enquiry_particulars.bid_opening_authority)
    _row(tbl, "Address", state.enquiry_particulars.address)
    _row(tbl, "Contact Details", state.enquiry_particulars.contact_details)
    _row(tbl, "Email", state.enquiry_particulars.email)

    # ─── Section 6: Bid Security ────────────────────────────────────
    doc.add_heading("6. Bid Security Details", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    _row(tbl, "Bid Security (INR)", f"Rs.{state.financial.bid_security_inr}.00")
    _row(tbl, "Bid Security In Favour Of", state.financial.bid_security_in_favour_of)
    _row(tbl, "Mode of Payment", state.financial.mode_of_payment)

    # ─── Section 7: Required Tender Documents ───────────────────────
    doc.add_heading("7. Required Tender Documents Details", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["S.No", "Document Name", "Stage", "Document Type"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for d in state.documents:
        row = tbl.add_row().cells
        row[0].text = str(d.s_no)
        row[1].text = d.document_name
        row[2].text = d.stage.value
        row[3].text = d.document_type.value

    # ─── Section 8: General Terms & Conditions / Eligibility ────────
    doc.add_heading("8. General Terms and Conditions / Eligibility", level=1)
    for para in state.general_terms.eligibility.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ─── Section 9: General Technical Terms ─────────────────────────
    doc.add_heading("9. General Technical Terms and Conditions (Procedure)", level=1)
    for para in state.general_terms.technical.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ─── Section 10: Legal Terms ────────────────────────────────────
    doc.add_heading("10. Legal Terms & Conditions", level=1)
    for para in state.general_terms.legal.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ─── Section 11: Procedure for Bid Submission ───────────────────
    doc.add_heading("11. Procedure for Bid Submission", level=1)
    for para in state.general_terms.bid_procedure.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ─── Section 12: Geographical Particulars ───────────────────────
    doc.add_heading("12. Geographical Particulars", level=1)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["S.No", "State", "District", "Mandal", "Assembly"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    row = tbl.add_row().cells
    row[0].text = "1"
    row[1].text = state.geography.state
    row[2].text = state.geography.district
    row[3].text = state.geography.mandal
    row[4].text = state.geography.assembly
    p = doc.add_paragraph()
    p.add_run(f"Parliament: {state.geography.parliament}").italic = True

    # ─── Section 13: Enquiry Forms (Stages) ─────────────────────────
    doc.add_heading("13. Enquiry Forms — Stage Details", level=1)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Stage", "Form Name", "Type of Form",
                           "Supporting Doc Required", "Supporting Doc Description"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for ef in state.enquiry_forms:
        row = tbl.add_row().cells
        row[0].text = ef.stage.value
        row[1].text = ef.form_name
        row[2].text = ef.type_of_form.value
        row[3].text = ef.supporting_document_required.value
        row[4].text = ef.supporting_document_description

    # ─── Section 14: BoQ summary (full BoQ is the XLSX artifact) ────
    doc.add_heading("14. Bill of Quantities (Summary)", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Total line items: {len(state.boq)}. Detailed BoQ available in BoQ.xlsx.").italic = True
    if state.boq:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["S.No", "Item", "Qty", "Unit"]):
            hdr[i].paragraphs[0].add_run(h).bold = True
        for r in state.boq:
            row = tbl.add_row().cells
            row[0].text = str(r.s_no)
            row[1].text = r.item
            row[2].text = str(r.qty)
            row[3].text = r.unit

    # ─── Section 15: Audit trail ────────────────────────────────────
    doc.add_heading("15. Audit Trail", level=1)
    doc.add_paragraph(f"Draft ID: {state.draft_id}")
    doc.add_paragraph(f"Version: v{state.version} ({state.current_gate.value})")
    doc.add_paragraph(f"Created by: {state.created_by} on {state.created_at}")
    doc.add_paragraph(f"Last updated: {state.last_updated_at}")
    if state.citations.rule_ids:
        doc.add_paragraph(f"Rules cited: {', '.join(state.citations.rule_ids)}").italic = True

    doc.save(out_path)
    return out_path


def _row(tbl, k: str, v: str) -> None:
    row = tbl.add_row().cells
    row[0].paragraphs[0].add_run(k).bold = True
    row[1].text = str(v) if v is not None else ""


# ─── PDF: same content via reportlab.platypus ───────────────────────


def render_bid_document_pdf(state: TenderDraftState, out_path: Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontSize=20,
        spaceAfter=8, textColor=colors.HexColor("#1F3864"),
    )
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=13,
                        spaceBefore=12, spaceAfter=6,
                        textColor=colors.HexColor("#1F3864"))
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=9,
                          spaceAfter=4, leading=12)

    story: list = []
    story.append(Paragraph("BID DOCUMENT", title_style))
    story.append(Paragraph(
        "<i>Government of Andhra Pradesh · eGP Portal</i>",
        body,
    ))
    story.append(Paragraph(
        f"Tender ID: <b>{state.tender_id or 'TBD'}</b>  ·  "
        f"NIT: <font face='Courier' size='8'>{state.tender_notice_number or ''}</font>",
        body,
    ))
    story.append(Spacer(1, 8))

    def kv_table(items: list[tuple[str, str]]) -> Table:
        rows = [[Paragraph(f"<b>{k}</b>", body), Paragraph(str(v or ""), body)] for k, v in items]
        tbl = Table(rows, colWidths=[55*mm, 110*mm])
        tbl.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9FB3D1")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EDF4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        return tbl

    # Same sections as DOCX but condensed for PDF
    story.append(Paragraph("1. Current Tender Details", h1))
    story.append(kv_table([
        ("Tender ID", state.tender_id or "(TBD)"),
        ("Name of Work", state.enquiry_particulars.name_of_work),
        ("Tender Category", state.classification.tender_category.value),
        ("Tender Type", state.classification.tender_type.value),
        ("Estimated Contract Value", f"₹ {state.financial.estimated_contract_value_inr:,}"),
        ("Submission Closing Date", state.dates.closing_date),
        ("Evaluation Type", state.evaluation.evaluation_type.value),
    ]))

    story.append(Paragraph("2. Enquiry Particulars", h1))
    story.append(kv_table([
        ("Department", state.enquiry_particulars.department_name),
        ("Circle / Division", state.enquiry_particulars.circle_division),
        ("Officer Inviting Bids", state.enquiry_particulars.officer_inviting_bids),
        ("Name of Project", state.enquiry_particulars.name_of_project),
        ("ECV (Words)", state.financial.estimated_contract_value_words),
        ("Period (months)", state.financial.period_of_completion_months),
        ("Form of Contract", state.classification.form_of_contract.value),
        ("Currency", f"{state.financial.currency_type.value} ({state.financial.default_currency})"),
    ]))

    story.append(Paragraph("3. Tender Dates", h1))
    story.append(kv_table([
        ("Start Date & Time", state.dates.start_date),
        ("End Date & Time", state.dates.end_date),
        ("Closing Date & Time", state.dates.closing_date),
        ("Bid Validity Period", f"{state.financial.bid_validity_days} days"),
    ]))

    story.append(Paragraph("4. Tender Inviting Authority", h1))
    story.append(kv_table([
        ("Officer", state.enquiry_particulars.officer_inviting_bids),
        ("Bid Opening Authority", state.enquiry_particulars.bid_opening_authority),
        ("Address", state.enquiry_particulars.address),
        ("Contact", state.enquiry_particulars.contact_details),
        ("Email", state.enquiry_particulars.email),
    ]))

    story.append(Paragraph("5. Bid Security Details", h1))
    story.append(kv_table([
        ("Bid Security (INR)", f"Rs.{state.financial.bid_security_inr}.00"),
        ("In Favour Of", state.financial.bid_security_in_favour_of),
        ("Mode of Payment", state.financial.mode_of_payment),
        ("Transaction Fee", f"₹{state.financial.transaction_fee_inr} ({state.financial.transaction_fee_payable_to})"),
    ]))

    story.append(PageBreak())

    story.append(Paragraph("6. Required Tender Documents", h1))
    doc_rows = [[Paragraph("<b>#</b>", body),
                 Paragraph("<b>Document</b>", body),
                 Paragraph("<b>Stage</b>", body),
                 Paragraph("<b>Type</b>", body)]]
    for d in state.documents:
        doc_rows.append([
            Paragraph(str(d.s_no), body),
            Paragraph(d.document_name, body),
            Paragraph(d.stage.value, body),
            Paragraph(d.document_type.value, body),
        ])
    tbl = Table(doc_rows, colWidths=[10*mm, 105*mm, 25*mm, 25*mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9FB3D1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(tbl)

    # Long-form sections — paragraph by paragraph
    for sect_title, sect_text in [
        ("7. General Terms & Conditions / Eligibility", state.general_terms.eligibility),
        ("8. General Technical Terms (Procedure)", state.general_terms.technical),
        ("9. Legal Terms & Conditions", state.general_terms.legal),
        ("10. Procedure for Bid Submission", state.general_terms.bid_procedure),
    ]:
        story.append(PageBreak())
        story.append(Paragraph(sect_title, h1))
        for para in sect_text.split("\n"):
            if para.strip():
                story.append(Paragraph(para.strip().replace("&", "&amp;"), body))

    story.append(PageBreak())
    story.append(Paragraph("11. Geographical Particulars", h1))
    story.append(kv_table([
        ("State", state.geography.state),
        ("District", state.geography.district),
        ("Mandal", state.geography.mandal),
        ("Assembly", state.geography.assembly),
        ("Parliament", state.geography.parliament),
    ]))

    story.append(Paragraph("12. Audit Trail", h1))
    story.append(kv_table([
        ("Draft ID", state.draft_id),
        ("Version", f"v{state.version} ({state.current_gate.value})"),
        ("Created", f"by {state.created_by} on {state.created_at}"),
        ("Last Updated", state.last_updated_at),
        ("Rules Cited", ", ".join(state.citations.rule_ids) or "(none)"),
    ]))

    pdf = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=15*mm, rightMargin=15*mm,
        topMargin=15*mm, bottomMargin=15*mm,
        title=f"BID DOCUMENT — {state.enquiry_particulars.name_of_work[:50]}",
        author="ProcureAI Module 1 (Drafter)",
    )

    def _footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(
            15*mm, 10*mm,
            f"Tender: {state.tender_id or state.draft_id}  ·  "
            f"v{state.version}  ·  Page {doc_.page}  ·  Generated: {state.last_updated_at}"
        )
        canvas.restoreState()

    pdf.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return out_path


# ─── BoQ as XLSX (or CSV fallback) ───────────────────────────────────


def render_boq_xlsx(state: TenderDraftState, out_path: Path) -> Path:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = "BoQ"
        # Header
        ws.cell(row=1, column=1, value=f"BILL OF QUANTITIES — {state.enquiry_particulars.name_of_work}")
        ws.cell(row=1, column=1).font = Font(bold=True, size=14)
        ws.merge_cells("A1:F1")

        ws.cell(row=2, column=1, value=f"Tender ID: {state.tender_id or '(TBD)'}")
        ws.cell(row=2, column=4, value=f"NIT: {state.tender_notice_number or ''}")

        # Column headers
        cols = ["S.No", "Item Description", "Qty", "Unit", "Rate (₹)", "Amount (₹)"]
        header_row = 4
        for i, c in enumerate(cols, start=1):
            cell = ws.cell(row=header_row, column=i, value=c)
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D9E2F3")
            cell.alignment = Alignment(horizontal="center")

        # Data
        for r, row in enumerate(state.boq, start=header_row + 1):
            ws.cell(row=r, column=1, value=row.s_no)
            ws.cell(row=r, column=2, value=row.item)
            ws.cell(row=r, column=3, value=row.qty)
            ws.cell(row=r, column=4, value=row.unit)
            ws.cell(row=r, column=5, value=row.rate or "")
            ws.cell(row=r, column=6, value=row.amount or "")

        # Footer
        total_row = header_row + len(state.boq) + 2
        ws.cell(row=total_row, column=5, value="GRAND TOTAL").font = Font(bold=True)
        ws.cell(row=total_row, column=6, value=f"= SUM(F{header_row+1}:F{total_row-2})").font = Font(bold=True)

        # Column widths
        for col, width in [("A", 6), ("B", 60), ("C", 10), ("D", 12), ("E", 14), ("F", 14)]:
            ws.column_dimensions[col].width = width

        wb.save(out_path)
        return out_path
    except ImportError:
        # CSV fallback
        csv_path = out_path.with_suffix(".csv")
        with open(csv_path, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["BoQ for", state.enquiry_particulars.name_of_work])
            writer.writerow(["Tender ID", state.tender_id or "(TBD)"])
            writer.writerow([])
            writer.writerow(["S.No", "Item", "Qty", "Unit", "Rate (₹)", "Amount (₹)"])
            for r in state.boq:
                writer.writerow([r.s_no, r.item, r.qty, r.unit, r.rate or "", r.amount or ""])
        return csv_path


# ─── Eligibility-only DOCX ──────────────────────────────────────────


def render_eligibility_docx(state: TenderDraftState, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("ELIGIBILITY CRITERIA", level=0)
    p = doc.add_paragraph()
    p.add_run(f"Tender: {state.enquiry_particulars.name_of_work}").bold = True
    doc.add_paragraph(f"Tender ID: {state.tender_id or '(TBD)'}")
    doc.add_paragraph(f"NIT: {state.tender_notice_number or ''}")
    doc.add_paragraph(
        f"ECV: ₹ {state.financial.estimated_contract_value_inr:,} "
        f"({state.financial.estimated_contract_value_words})"
    )
    doc.add_paragraph("")

    for para in state.general_terms.eligibility.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Cited rules
    if state.citations.rule_ids:
        doc.add_heading("Source Rules Cited", level=1)
        for s in state.citations.sources:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{s.node_id}: ").bold = True
            p.add_run(s.quote_excerpt).italic = True

    doc.save(out_path)
    return out_path


# ─── Markdown summary (convenience artifact) ────────────────────────


def render_md_summary(state: TenderDraftState, out_path: Path) -> Path:
    lines: list[str] = []
    ep = state.enquiry_particulars
    f = state.financial
    g = state.geography
    lines.append(f"# Bid Document — {ep.name_of_work}\n")
    lines.append(f"**Tender ID:** {state.tender_id or '(TBD)'}  ·  **NIT:** {state.tender_notice_number or ''}")
    lines.append(f"**Version:** v{state.version} ({state.current_gate.value})  ·  **Last updated:** {state.last_updated_at}\n")
    lines.append("---\n")

    lines.append("## Authority")
    lines.append(f"- **Department:** {ep.department_name}")
    lines.append(f"- **Circle / Division:** {ep.circle_division}")
    lines.append(f"- **Officer:** {ep.officer_inviting_bids}")
    lines.append(f"- **Contact:** {ep.email} · {ep.contact_details}")
    lines.append(f"- **Address:** {ep.address}\n")

    lines.append("## Tender")
    lines.append(f"- **Category:** {state.classification.tender_category.value}  ·  **Type:** {state.classification.tender_type.value}  ·  **Form:** {state.classification.form_of_contract.value}")
    lines.append(f"- **ECV:** ₹ {f.estimated_contract_value_inr:,} ({f.estimated_contract_value_words})")
    lines.append(f"- **Period:** {f.period_of_completion_months} months  ·  **Bid Validity:** {f.bid_validity_days} days")
    lines.append(f"- **Bid Security:** {f.bid_security_percent}% = ₹{f.bid_security_inr:,}\n")

    lines.append("## Geography")
    lines.append(f"{g.state} → {g.district} → {g.mandal} ({g.assembly}) · Parliament: {g.parliament}\n")

    lines.append("## Eligibility")
    lines.append("```")
    lines.append(state.general_terms.eligibility)
    lines.append("```\n")

    lines.append("## Technical Terms")
    lines.append("```")
    lines.append(state.general_terms.technical)
    lines.append("```\n")

    lines.append("## Legal Terms")
    lines.append("```")
    lines.append(state.general_terms.legal)
    lines.append("```\n")

    lines.append("## Bid Procedure")
    lines.append("```")
    lines.append(state.general_terms.bid_procedure)
    lines.append("```\n")

    lines.append(f"## BoQ ({len(state.boq)} items)\n")
    lines.append("| # | Item | Qty | Unit |")
    lines.append("|---|---|---:|---|")
    for r in state.boq:
        lines.append(f"| {r.s_no} | {r.item} | {r.qty} | {r.unit} |")
    lines.append("")

    lines.append("## Rules Cited")
    for s in state.citations.sources:
        lines.append(f"- **{s.node_id}** — {s.quote_excerpt}")

    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


# ─── Orchestrator: render all artifacts on publish ──────────────────


def render_all_for_publish(state: TenderDraftState) -> dict:
    out_dir = _artifact_dir(state)
    artifacts: dict[str, str] = {}

    bid_docx = out_dir / "BID_DOCUMENT.docx"
    bid_pdf = out_dir / "BID_DOCUMENT.pdf"
    boq_xlsx = out_dir / "BoQ.xlsx"
    elig_docx = out_dir / "ELIGIBILITY.docx"
    md_summary = out_dir / "summary.md"

    try:
        render_bid_document_docx(state, bid_docx)
        artifacts["bid_document_docx"] = str(bid_docx)
    except Exception as e:
        artifacts["bid_document_docx_error"] = str(e)

    try:
        render_bid_document_pdf(state, bid_pdf)
        artifacts["bid_document_pdf"] = str(bid_pdf)
    except Exception as e:
        artifacts["bid_document_pdf_error"] = str(e)

    try:
        boq_path = render_boq_xlsx(state, boq_xlsx)
        artifacts["boq_xlsx"] = str(boq_path)
    except Exception as e:
        artifacts["boq_xlsx_error"] = str(e)

    try:
        render_eligibility_docx(state, elig_docx)
        artifacts["eligibility_docx"] = str(elig_docx)
    except Exception as e:
        artifacts["eligibility_docx_error"] = str(e)

    try:
        render_md_summary(state, md_summary)
        artifacts["summary_md"] = str(md_summary)
    except Exception as e:
        artifacts["summary_md_error"] = str(e)

    return {"artifact_dir": str(out_dir), **artifacts}
